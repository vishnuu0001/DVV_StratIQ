# ---------------------------------------------------------------------------
# Author: Vishnuu A
# Scope: Shared python-docx rendering for BRD/FSD/Solution Documentation/Test Plan (spec §5
# Date: 2025-12-15
# ---------------------------------------------------------------------------
"""Shared python-docx rendering for BRD/FSD/Solution Documentation/Test Plan (spec §5
Agent 2's implementation notes, generalised to every document-generating agent this
pass adds).

Two modes:
- **Template mode** (`Project.brd_template_id` set): open the client's `.dotx`, locate
  `{{SECTION:key}}` placeholder paragraphs, replace them in-place at the XML level
  (preserving surrounding formatting), reference the template's own named styles
  rather than building new ones.
- **Default mode** (no template): a built-in style set so the pipeline works out of
  the box without requiring a pre-uploaded template.

Either way: regenerate the TOC field so page numbers resolve on first open, and stamp
document properties.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_PLACEHOLDER_RE = re.compile(r"\{\{SECTION:(\w+)\}\}")


@dataclass
class RenderedSection:
    key: str
    heading: str
    mode: str  # GENERATED | REQ_TABLE | GLOSSARY | RTM_SUMMARY
    body_text: str | None = None  # GENERATED prose (already LLM-authored by the caller)
    citations: list[str] = field(default_factory=list)  # footnote strings, GENERATED mode
    table_rows: list[dict] = field(default_factory=list)  # REQ_TABLE / GLOSSARY / RTM_SUMMARY
    table_columns: list[str] = field(default_factory=list)


# Function: render_document
def render_document(
    *, template_path: str | None, title: str, subtitle: str, sections: list[RenderedSection],
    output_path: str, author: str = "TraceForge",
) -> None:
    if template_path:
        doc = Document(template_path)
        _render_into_template(doc, sections)
    else:
        doc = Document()
        _apply_default_styles(doc)
        _render_default(doc, title, subtitle, sections)

    _set_core_properties(doc, title, author)
    _insert_or_refresh_toc(doc)
    doc.save(output_path)


# ── Default (no client template) rendering ───────────────────────

# Function: _apply_default_styles
def _apply_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


# Function: _render_default
def _render_default(doc: Document, title: str, subtitle: str, sections: list[RenderedSection]) -> None:
    doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    doc.add_heading("Table of Contents", level=1)
    _add_toc_field(doc)
    doc.add_page_break()

    for section in sections:
        doc.add_heading(section.heading, level=1)
        _render_section_body(doc, section)


# Function: _render_section_body
def _render_section_body(doc: Document, section: RenderedSection) -> None:
    if section.mode == "GENERATED":
        for paragraph_text in (section.body_text or "").split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
            doc.add_paragraph(paragraph_text)
        if section.citations:
            note = doc.add_paragraph()
            note.add_run("Sources: " + "; ".join(section.citations)).italic = True
    elif section.mode in ("REQ_TABLE", "GLOSSARY", "RTM_SUMMARY"):
        _render_table(doc, section.table_columns, section.table_rows)
    else:
        raise ValueError(f"Unknown section mode: {section.mode}")


# Function: _render_table
def _render_table(doc: Document, columns: list[str], rows: list[dict]) -> None:
    if not rows:
        doc.add_paragraph("(none)")
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid" if "Table Grid" in [s.name for s in doc.styles] else None
    header = table.rows[0].cells
    for i, col in enumerate(columns):
        header[i].text = col
        for run in header[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row.get(col, ""))


# ── Template mode: {{SECTION:key}} placeholder replacement ──────

# Function: _render_into_template
def _render_into_template(doc: Document, sections: list[RenderedSection]) -> None:
    by_key = {s.key: s for s in sections}
    for paragraph in list(doc.paragraphs):
        match = _PLACEHOLDER_RE.search(paragraph.text)
        if not match:
            continue
        key = match.group(1)
        section = by_key.get(key)
        anchor = paragraph._element
        if section is None:
            anchor.getparent().remove(anchor)
            continue
        _insert_section_at(doc, anchor, section)
        anchor.getparent().remove(anchor)  # remove the placeholder itself, after inserting content before it


# Function: _insert_generated_section
def _insert_generated_section(doc: Document, anchor_element, section: RenderedSection) -> None:
    for paragraph_text in (section.body_text or "").split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        p = doc.add_paragraph(paragraph_text)
        anchor_element.addprevious(p._element)
    if section.citations:
        note = doc.add_paragraph("Sources: " + "; ".join(section.citations))
        note.runs[0].italic = True
        anchor_element.addprevious(note._element)


# Function: _insert_table_section
def _insert_table_section(doc: Document, anchor_element, section: RenderedSection) -> None:
    table_doc_position = doc.add_table(rows=1, cols=max(1, len(section.table_columns)))
    anchor_element.addprevious(table_doc_position._element)
    style_name = "Table Grid" if "Table Grid" in [s.name for s in doc.styles] else None
    if style_name:
        table_doc_position.style = style_name
    header = table_doc_position.rows[0].cells
    for i, col in enumerate(section.table_columns):
        header[i].text = col
    for row in section.table_rows:
        cells = table_doc_position.add_row().cells
        for i, col in enumerate(section.table_columns):
            cells[i].text = str(row.get(col, ""))


# Function: _insert_section_at
def _insert_section_at(doc: Document, anchor_element, section: RenderedSection) -> None:
    """Insert new paragraphs/tables immediately before `anchor_element`, preserving
    the template's surrounding formatting (per spec: XML-level in-place replacement,
    not a full-document rebuild)."""
    heading = doc.add_paragraph(section.heading, style=_first_available_style(doc, ["Heading 1", "Heading1"]))
    anchor_element.addprevious(heading._element)

    if section.mode == "GENERATED":
        _insert_generated_section(doc, anchor_element, section)
    else:
        _insert_table_section(doc, anchor_element, section)


# Function: _first_available_style
def _first_available_style(doc: Document, candidates: list[str]) -> str | None:
    names = {s.name for s in doc.styles}
    for candidate in candidates:
        if candidate in names:
            return candidate
    return None  # fail loudly at template-registration time is the spec's ask; here we degrade to default style


# ── TOC + document properties ────────────────────────────────────

# Function: _add_toc_field
def _add_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    field_char_begin = OxmlElement("w:fldChar")
    field_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    field_char_separate = OxmlElement("w:fldChar")
    field_char_separate.set(qn("w:fldCharType"), "separate")
    field_char_end = OxmlElement("w:fldChar")
    field_char_end.set(qn("w:fldCharType"), "end")
    for element in (field_char_begin, instr_text, field_char_separate, field_char_end):
        run._r.append(element)


# Function: _insert_or_refresh_toc
def _insert_or_refresh_toc(doc: Document) -> None:
    """Sets updateFields in settings.xml so Word recalculates the TOC/page numbers on
    first open (spec §5's exact instruction) — python-docx can't compute real page
    numbers itself (that requires a layout engine), so this defers that to Word."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


# Function: _set_core_properties
def _set_core_properties(doc: Document, title: str, author: str) -> None:
    props = doc.core_properties
    props.title = title
    props.author = author
    props.created = datetime.now(timezone.utc)
    props.subject = "Generated by TraceForge"
