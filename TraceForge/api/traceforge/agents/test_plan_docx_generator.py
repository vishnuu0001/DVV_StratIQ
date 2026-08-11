# ---------------------------------------------------------------------------
# Author: TraceForge Team
# Scope: Rich TestPlan DOCX generation with comprehensive documentation structure
# Date: 2026-02-20
# ---------------------------------------------------------------------------
"""Rich TestPlan DOCX generation matching reference quality standards.

Generates comprehensive test plan documents with:
- Executive Summary
- Source Requirement Baseline
- Test Objectives
- Scope (In/Out)
- Test Strategy & Approach
- Test Environment & Infrastructure  
- Schedule & Phases
- Entry/Exit Criteria
- Test Data Requirements
- Quality Metrics & Success Criteria
- Detailed tables with test coverage metadata
"""
from __future__ import annotations

import json
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from traceforge.agents.base import call_agent_llm
from traceforge.config import FAST_PIPELINE, STORAGE_DIR
from traceforge.db.models import Artifact, Project, Requirement, TestCase, TestPlan
from traceforge.llm.ollama import OllamaProvider


# Function: _add_heading
def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(12)
            run.font.bold = True


# Function: _add_section_text
def _add_section_text(doc: Document, text: str) -> None:
    """Add body text with proper formatting."""
    if not text:
        return
    for paragraph_text in text.split("\n\n"):
        para_text = paragraph_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.line_spacing = 1.15


# Function: _add_bullet_list
def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bulleted list."""
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


# Function: _add_table_from_rows
def _add_table_from_rows(doc: Document, headers: list[str], rows: list[dict]) -> None:
    """Add a table to the document."""
    if not rows:
        doc.add_paragraph("(No data available)")
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    
    # Format header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i]._element.get_or_add_tcPr().append(
            _get_shading_elm("1D4ED8")  # Dark blue background
        )
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row_data.get(header, ""))


def _test_case_metadata(test_case: TestCase) -> dict:
    raw = test_case.gherkin or ""
    if not raw.lstrip().startswith("{"):
        return {}
    try:
        parsed = json.loads(raw)
    except (TypeError, ValueError):
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _regression_coverage_rows(requirements: list[Requirement], test_cases: list[TestCase]) -> list[dict]:
    cases_by_requirement: dict[uuid.UUID, list[TestCase]] = {}
    for test_case in test_cases:
        cases_by_requirement.setdefault(test_case.requirement_id, []).append(test_case)
    rows = []
    for requirement in requirements:
        cases = cases_by_requirement.get(requirement.id, [])
        rows.append({
            "REQ-ID": requirement.req_id,
            "Requirement": requirement.title,
            "Scenarios": len(cases),
            "P1/P2/P3": "/".join(str(sum(case.priority == priority for case in cases)) for priority in ("P1", "P2", "P3")),
            "Test Types": ", ".join(sorted({case.test_type for case in cases})) or "COVERAGE GAP",
            "UI Candidates": sum(case.test_level == "UI_E2E" for case in cases),
        })
    return rows


# Function: _get_shading_elm
def _get_shading_elm(color_hex: str):
    """Create a shading element for cell background."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    return shading


# Function: _generate_executive_summary
async def _generate_executive_summary(
    session: AsyncSession,
    provider: OllamaProvider,
    project: Project,
    requirements: list[Requirement],
    pipeline_run_id: uuid.UUID | None,
) -> str:
    """Generate executive summary section."""
    if FAST_PIPELINE:
        return (
            f"This comprehensive test plan validates all {len(requirements)} approved requirements "
            f"for the {project.name} project. The plan covers end-to-end business processes, integration points, "
            f"and quality gates. Testing is risk-based, requirement-driven, and structured for execution "
            f"across multiple environments and phases."
        )
    
    req_summary = "\n".join(f"- {r.req_id}: {r.statement}" for r in requirements[:20])
    system = (
        "You are a senior test lead. Write a 150-200 word executive summary for a test plan. "
        "Ground it in the provided requirements. Return plain prose only (no JSON, no markdown)."
    )
    user = f"Project: {project.name}\nClient: {project.client_name or 'TBD'}\n\nRequirements:\n{req_summary}"
    
    response = await provider.generate(system, user, temperature=0.3, max_tokens=400, json_mode=False)
    return response.text.strip()


# Function: _generate_test_objectives
async def _generate_test_objectives(
    session: AsyncSession,
    provider: OllamaProvider,
    project: Project,
    requirements: list[Requirement],
    pipeline_run_id: uuid.UUID | None,
) -> list[str]:
    """Generate 8+ detailed test objectives from requirements."""
    if FAST_PIPELINE:
        objectives = [
            "Validate all approved business requirements are correctly implemented and functioning",
            "Verify functional requirements are met with correct data flow and processing logic",
            "Validate non-functional requirements including performance, security, and scalability",
            "Test integration between all internal and external systems",
            "Verify data integrity, consistency, and persistence across system components",
            "Validate user workflows and business process completeness",
            "Test error handling, edge cases, and negative scenarios",
            "Verify compliance with regulatory and security requirements",
        ]
        return objectives[:8]
    
    req_text = "\n".join(f"- {r.req_id} [{r.level}]: {r.statement}" for r in requirements[:30])
    system = (
        "You are a test architect. Generate 8-10 specific, measurable test objectives "
        "from the requirements below. Each objective should start with 'Prove that' or 'Verify that' "
        "and be testable. Return only a JSON array of strings, nothing else: "
        '{"objectives": ["Objective 1", "Objective 2", ...]}'
    )
    user = f"Requirements:\n{req_text}"
    
    response = await provider.generate(system, user, temperature=0.3, max_tokens=600, json_mode=True)
    try:
        import json
        parsed = json.loads(response.text)
        return parsed.get("objectives", [])[:10]
    except:
        return []


# Function: _generate_test_strategy
async def _generate_test_strategy(
    session: AsyncSession,
    provider: OllamaProvider,
    project: Project,
    requirements: list[Requirement],
    pipeline_run_id: uuid.UUID | None,
) -> str:
    """Generate test strategy section."""
    if FAST_PIPELINE:
        return (
            "This plan employs a risk-based testing strategy with requirement-level traceability. "
            "Testing is organized in phases: (1) Unit/Component validation, (2) Integration testing, "
            "(3) End-to-End scenarios, (4) Regression and UAT. Each test case maps to one or more requirements "
            "and is classified by type (POSITIVE, NEGATIVE, EDGE, BOUNDARY, SECURITY, PERFORMANCE) and priority (P1/P2/P3). "
            "Test data is production-realistic and sourced from business stakeholders. All test cases are automated via Playwright "
            "unless a specific scenario requires manual validation."
        )
    
    req_summary = "\n".join(f"- {r.req_id} [{r.level}]: {r.statement}" for r in requirements[:20])
    system = (
        "You are a test architect. Write a 200-300 word test strategy section for a comprehensive test plan. "
        "Include: testing phases, risk-based approach, test types, automation approach, test data strategy. "
        "Return plain prose only (no JSON, no markdown)."
    )
    user = f"Project: {project.name}\nRequirements:\n{req_summary}"
    
    response = await provider.generate(system, user, temperature=0.3, max_tokens=600, json_mode=False)
    return response.text.strip()


# Function: _generate_test_scope
async def _generate_test_scope(
    session: AsyncSession,
    provider: OllamaProvider,
    project: Project,
    requirements: list[Requirement],
    pipeline_run_id: uuid.UUID | None,
) -> tuple[list[str], list[str]]:
    """Generate In-Scope and Out-of-Scope items."""
    if FAST_PIPELINE:
        in_scope = [
            f"All {len(requirements)} APPROVED requirements",
            "End-to-end business process flows",
            "Integration with external systems and APIs",
            "Data validation and business rules",
            "Performance under normal load conditions",
            "Security and access control validations",
        ]
        out_of_scope = [
            "Load testing beyond normal peak volume",
            "Third-party application internal testing",
            "Operational deployment procedures",
            "Documentation beyond test artifacts",
        ]
        return in_scope, out_of_scope
    
    req_text = "\n".join(f"- {r.req_id}: {r.statement}" for r in requirements[:20])
    system = (
        "You are a test architect. Based on the requirements below, generate 6-8 in-scope "
        "and 4-6 out-of-scope items for a test plan. Return JSON: "
        '{"in_scope": [...], "out_of_scope": [...]}'
    )
    user = f"Requirements:\n{req_text}"
    
    response = await provider.generate(system, user, temperature=0.3, max_tokens=600, json_mode=True)
    try:
        import json
        parsed = json.loads(response.text)
        return parsed.get("in_scope", []), parsed.get("out_of_scope", [])
    except:
        return [], []


async def generate_test_plan_docx(
    session: AsyncSession,
    *,
    project_id: uuid.UUID,
    test_plan: TestPlan,
    pipeline_run_id: uuid.UUID | None,
) -> str:
    """Generate a comprehensive test plan DOCX document.
    
    Returns the path to the generated DOCX file.
    """
    project = await session.get(Project, project_id)
    if not project:
        raise ValueError(f"Project {project_id} not found")
    
    # Fetch all approved requirements and test cases
    result = await session.execute(
        select(Requirement).where(
            Requirement.project_id == project_id,
            Requirement.status == "APPROVED"
        ).order_by(Requirement.req_id)
    )
    requirements = list(result.scalars().all())
    
    result = await session.execute(
        select(TestCase).where(TestCase.project_id == project_id)
        .order_by(TestCase.test_type, TestCase.priority)
    )
    test_cases = list(result.scalars().all())
    
    # Initialize LLM provider
    provider = OllamaProvider()
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading(project.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("DETAILED END-TO-END TEST PLAN")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    doc.add_paragraph()
    
    metadata_table = doc.add_table(rows=5, cols=2)
    metadata_table.style = "Light Grid Accent 1"
    metadata = [
        ("Scenario ID", f"{project.key}"),
        ("Wave / Release", test_plan.schedule.get("phases", ["TBD"])[0] if test_plan.schedule else "TBD"),
        ("Mill(s) / Site(s) Involved", project.config.get("site", "TBD")),
        ("Client", project.client_name or "TBD"),
        ("Test Plan Version", str(test_plan.version)),
    ]
    for i, (key, value) in enumerate(metadata):
        metadata_table.rows[i].cells[0].text = key
        metadata_table.rows[i].cells[1].text = value
        for run in metadata_table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
    
    doc.add_page_break()
    
    # ===== 1. EXECUTIVE SUMMARY =====
    _add_heading(doc, "1. Executive Summary", level=1)
    exec_summary = test_plan.scope
    _add_section_text(doc, exec_summary)
    
    # Add companion Excel note
    doc.add_paragraph(
        f"The companion Excel workbook contains {len(test_cases)} DRAFT test cases "
        f"with requirement mapping and reviewed generation metadata."
    )
    
    # ===== 2. SOURCE REQUIREMENT BASELINE =====
    _add_heading(doc, "2. Source Requirement Baseline", level=1)
    doc.add_paragraph(
        f"This test plan is based on {len(requirements)} APPROVED requirements across all levels "
        "(Business, Functional, Non-Functional, Constraints, Assumptions)."
    )
    
    req_table_data = []
    for req in requirements[:30]:  # Show first 30
        req_table_data.append({
            "REQ-ID": req.req_id,
            "Level": req.level,
            "Status": req.status,
            "Priority": req.priority,
        })
    _add_table_from_rows(
        doc,
        ["REQ-ID", "Level", "Status", "Priority"],
        req_table_data,
    )
    
    # ===== 3. TEST OBJECTIVES =====
    _add_heading(doc, "3. Test Objectives", level=1)
    objectives = [
        f"{req.req_id}: {criterion}"
        for req in requirements
        for criterion in (req.acceptance_criteria or [req.statement])
    ]
    _add_bullet_list(doc, objectives or ["PENDING BUSINESS CONFIRMATION"])
    
    # ===== 4. SCOPE =====
    _add_heading(doc, "4. Scope", level=1)
    
    schedule = test_plan.schedule or {}
    in_scope = schedule.get("in_scope") or [f"{req.req_id}: {req.statement}" for req in requirements]
    out_of_scope = schedule.get("out_of_scope") or ["PENDING BUSINESS CONFIRMATION"]
    
    _add_heading(doc, "4.1 In Scope", level=2)
    _add_bullet_list(doc, in_scope or ["PENDING BUSINESS CONFIRMATION"])
    
    _add_heading(doc, "4.2 Out of Scope", level=2)
    _add_bullet_list(doc, out_of_scope)
    
    # ===== 5. TEST STRATEGY & APPROACH =====
    _add_heading(doc, "5. Test Strategy & Approach", level=1)
    strategy = test_plan.strategy
    _add_section_text(doc, strategy)
    
    # ===== 6. TEST ENVIRONMENT & INFRASTRUCTURE =====
    _add_heading(doc, "6. Test Environment & Infrastructure", level=1)
    _add_heading(doc, "6.1 Environments", level=2)
    _add_bullet_list(doc, test_plan.environments or ["PENDING BUSINESS CONFIRMATION"])
    
    _add_heading(doc, "6.2 Infrastructure Requirements", level=2)
    _add_bullet_list(doc, ["PENDING BUSINESS CONFIRMATION"])
    
    # ===== 7. TEST SCHEDULE & PHASES =====
    _add_heading(doc, "7. Test Schedule & Phases", level=1)
    phases = schedule.get("phases") or ["PENDING BUSINESS CONFIRMATION"]
    _add_bullet_list(doc, phases)
    for heading, key in (
        ("7.1 Execution Waves", "execution_waves"),
        ("7.2 Process Coverage", "process_stages"),
        ("7.3 Test Data Strategy", "test_data_strategy"),
        ("7.4 Role and Access Strategy", "role_strategy"),
        ("7.5 Automation Strategy", "automation_strategy"),
        ("7.6 Risks and Dependencies", "risks"),
        ("7.7 Deliverables", "deliverables"),
    ):
        _add_heading(doc, heading, level=2)
        _add_bullet_list(doc, schedule.get(key) or ["PENDING BUSINESS CONFIRMATION"])
    
    # ===== 8. ENTRY & EXIT CRITERIA =====
    _add_heading(doc, "8. Entry & Exit Criteria", level=1)
    
    entry_exit = test_plan.entry_exit_criteria or {}
    entry_criteria = entry_exit.get("entry") or ["PENDING BUSINESS CONFIRMATION"]
    exit_criteria = entry_exit.get("exit") or ["PENDING BUSINESS CONFIRMATION"]
    suspension_criteria = entry_exit.get("suspension") or ["PENDING BUSINESS CONFIRMATION"]
    resumption_criteria = entry_exit.get("resumption") or ["PENDING BUSINESS CONFIRMATION"]
    
    _add_heading(doc, "8.1 Entry Criteria", level=2)
    _add_bullet_list(doc, entry_criteria)
    
    _add_heading(doc, "8.2 Exit Criteria", level=2)
    _add_bullet_list(doc, exit_criteria)

    _add_heading(doc, "8.3 Suspension Criteria", level=2)
    _add_bullet_list(doc, suspension_criteria)

    _add_heading(doc, "8.4 Resumption Criteria", level=2)
    _add_bullet_list(doc, resumption_criteria)
    
    # ===== 9. TEST CASE MATRIX =====
    _add_heading(doc, "9. Test Case Coverage Matrix", level=1)
    
    # Count by type
    type_counts = {}
    for tc in test_cases:
        type_counts[tc.test_type] = type_counts.get(tc.test_type, 0) + 1
    
    matrix_data = [
        {"Test Type": "POSITIVE", "Count": type_counts.get("POSITIVE", 0), "Description": "Happy path and main flow scenarios"},
        {"Test Type": "NEGATIVE", "Count": type_counts.get("NEGATIVE", 0), "Description": "Error handling and business rule rejection"},
        {"Test Type": "EDGE", "Count": type_counts.get("EDGE", 0), "Description": "Boundary, retry, and recovery scenarios"},
        {"Test Type": "BOUNDARY", "Count": type_counts.get("BOUNDARY", 0), "Description": "Min/max and limit testing"},
        {"Test Type": "NEGATIVE_SECURITY", "Count": type_counts.get("NEGATIVE_SECURITY", 0), "Description": "Security and access control"},
        {"Test Type": "PERFORMANCE", "Count": type_counts.get("PERFORMANCE", 0), "Description": "Load and performance testing"},
    ]
    _add_table_from_rows(
        doc,
        ["Test Type", "Count", "Description"],
        matrix_data,
    )
    
    doc.add_paragraph(f"Total Test Cases: {len(test_cases)}")
    
    # ===== 10. REQUIREMENT COVERAGE & REGRESSION SCOPE =====
    _add_heading(doc, "10. Requirement Coverage & Automation Regression Scope", level=1)
    doc.add_paragraph(
        "The full regression baseline includes every scenario below. P1 scenarios form the critical-path "
        "regression subset; P2 and P3 scenarios extend business-rule, boundary, recovery, integration, and "
        "non-functional coverage where supported by the approved requirement evidence. UI E2E scenarios may "
        "be generated as skipped Playwright placeholders until environment bindings are configured."
    )
    coverage_rows = _regression_coverage_rows(requirements, test_cases)
    _add_table_from_rows(
        doc,
        ["REQ-ID", "Requirement", "Scenarios", "P1/P2/P3", "Test Types", "UI Candidates"],
        coverage_rows,
    )
    uncovered = [row["REQ-ID"] for row in coverage_rows if row["Scenarios"] == 0]
    doc.add_paragraph(
        "Requirement coverage gaps: " + (", ".join(uncovered) if uncovered else "None")
    )

    # ===== 11. DETAILED REGRESSION SCENARIOS =====
    _add_heading(doc, "11. Detailed Automation Regression Scenarios", level=1)
    requirement_by_id = {requirement.id: requirement for requirement in requirements}
    ordered_cases = sorted(
        test_cases,
        key=lambda case: (requirement_by_id.get(case.requirement_id).req_id if case.requirement_id in requirement_by_id else "", case.tc_id),
    )
    for test_case in ordered_cases:
        requirement = requirement_by_id.get(test_case.requirement_id)
        metadata = _test_case_metadata(test_case)
        automation_status = metadata.get("automation_status", "AUTOMATION_BLOCKED")
        _add_heading(
            doc,
            f"{test_case.tc_id}: {test_case.title}",
            level=2,
        )
        doc.add_paragraph(
            f"Requirement: {requirement.req_id if requirement else 'UNMAPPED'} | "
            f"Type: {test_case.test_type} | Level: {test_case.test_level} | "
            f"Priority: {test_case.priority} | Risk: {metadata.get('risk_rating', 'Not classified')} | "
            f"Status: {test_case.status} | Automation: {automation_status}"
        )
        doc.add_paragraph(f"Objective: {metadata.get('objective') or test_case.title}")
        _add_heading(doc, "Preconditions", level=3)
        _add_bullet_list(doc, test_case.preconditions or ["No source-confirmed preconditions recorded"])
        _add_heading(doc, "Execution Steps and Expected Results", level=3)
        for index, step in enumerate(test_case.steps or [], start=1):
            step_number = step.get("step_no", index)
            doc.add_paragraph(
                f"Step {step_number} - Action: {step.get('action', '')}\n"
                f"Test Data: {step.get('test_data', '') or 'No source-confirmed data recorded'}\n"
                f"Expected Result: {step.get('expected_result', '')}"
            )
        if not test_case.steps:
            doc.add_paragraph("No executable steps were generated; business review is required.")
        blockers = metadata.get("automation_blockers") or []
        if blockers:
            doc.add_paragraph("Automation configuration required: " + "; ".join(map(str, blockers)))

    # ===== 12. SUCCESS CRITERIA & METRICS =====
    _add_heading(doc, "12. Quality Metrics & Success Criteria", level=1)
    metrics_data = [
        {"Metric": "Approved requirement count", "Target": "Informational", "Measurement": str(len(requirements))},
        {"Metric": "Draft test-case count", "Target": "Informational", "Measurement": str(len(test_cases))},
        {"Metric": "Execution and defect thresholds", "Target": "PENDING BUSINESS CONFIRMATION", "Measurement": "PENDING BUSINESS CONFIRMATION"},
    ]
    _add_table_from_rows(
        doc,
        ["Metric", "Target", "Measurement"],
        metrics_data,
    )
    
    # ===== FOOTER METADATA =====
    doc.add_page_break()
    _add_heading(doc, "Document Information", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Test Plan Version: {test_plan.version}")
    doc.add_paragraph(f"Status: {test_plan.status}")
    doc.add_paragraph(f"Total Requirements: {len(requirements)}")
    doc.add_paragraph(f"Total Test Cases: {len(test_cases)}")
    
    # Save document
    project_dir = STORAGE_DIR / str(project_id) / "artifacts"
    project_dir.mkdir(parents=True, exist_ok=True)
    
    filename = f"{project.key}_Test_Plan_v{test_plan.version}.docx"
    output_path = project_dir / filename
    
    doc.save(str(output_path))
    
    return str(output_path)
