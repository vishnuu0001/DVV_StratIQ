#!/usr/bin/env python3
import json
import sys

# Read requirement from docx
try:
    from docx import Document
    doc = Document(r"C:\RD\DVV_StratIQ-Aqorynth\dsvstratiq\TraceForge\data\TestData\E2E_OB_Direct Order – Twin Reel With FSC Requirements.docx")
    print("=== REQUIREMENT DOCUMENT ===")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"{i}: {para.text[:150]}")
except Exception as e:
    print(f"Error reading docx: {e}")

# Read test plan from docx
try:
    doc2 = Document(r"C:\RD\DVV_StratIQ-Aqorynth\dsvstratiq\TraceForge\data\TestData\Outbound_E2E_01_Detailed_Test_Plan.docx")
    print("\n=== TEST PLAN DOCUMENT ===")
    for i, para in enumerate(doc2.paragraphs[:25]):
        if para.text.strip():
            print(f"{i}: {para.text[:150]}")
except Exception as e:
    print(f"Error reading test plan: {e}")

# Read test cases from xlsx
try:
    from openpyxl import load_workbook
    wb = load_workbook(r"C:\RD\DVV_StratIQ-Aqorynth\dsvstratiq\TraceForge\data\TestData\Outbound_E2E_01_Detailed_Test_Cases.xlsx")
    ws = wb.active
    print(f"\n=== TEST CASES SPREADSHEET ({ws.title}) ===")
    print(f"Dimensions: {ws.dimensions}")
    print("\nHeaders and first 15 rows:")
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=16, values_only=True)):
        print(f"Row {i}: {row[:8]}")
    
    # Count total test cases
    row_count = ws.max_row
    print(f"\nTotal rows (including header): {row_count}")
    print(f"Estimated test cases: {row_count - 1}")
except Exception as e:
    print(f"Error reading xlsx: {e}")
