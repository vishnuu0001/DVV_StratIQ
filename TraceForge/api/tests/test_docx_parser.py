from docx import Document

from traceforge.parsing.docx import parse_docx


def test_docx_parser_preserves_interleaved_table_order_and_cell_relationships(tmp_path):
    path = tmp_path / "interleaved.docx"
    document = Document()
    document.add_heading("References", level=1)
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Grade\nG-100"
    table.cell(0, 1).text = "Customer\nExample Ltd"
    document.add_paragraph("After table")
    document.save(path)

    parsed = parse_docx(str(path))

    assert [block.text for block in parsed.blocks] == [
        "References",
        "Before table",
        "Grade G-100 | Customer Example Ltd",
        "After table",
    ]
    assert all(block.section_path == "References" for block in parsed.blocks)
