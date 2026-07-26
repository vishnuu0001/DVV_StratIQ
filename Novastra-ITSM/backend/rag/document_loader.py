# ---------------------------------------------------------------------------
# Author: Vishnuu A
# Scope: Document loading utilities.
# Date: 2025-09-16
# ---------------------------------------------------------------------------
"""
Document loading utilities.
Supports: .docx, .xlsx, .csv, .txt, .md (structured), .pdf, .png/.jpg (OCR).

Chunking strategy:
  - DOCX / XLSX / MD : structure-aware → one atomic Document per incident
    (Issue + Analysis + Solution kept together; splitter never breaks them)
  - TXT / CSV / PDF  : generic RecursiveCharacterTextSplitter after load
  - Images           : OCR → single Document per file

All Documents carry enriched metadata (incident_number, solman_id,
delivery_numbers, oss_notes) enabling exact-ID lookups before semantic search.
"""
import io
import logging
import re
from pathlib import Path
from typing import List

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# ── SAP entity extraction ─────────────────────────────────────────────────
_SAP_INC_RE      = re.compile(r'\b(INC\d+)\b', re.IGNORECASE)
_SAP_SOLMAN_RE   = re.compile(r'SolManID#?\s*(\d{10,13})', re.IGNORECASE)
_SAP_DELIVERY_RE = re.compile(
    r'(?:delivery|transfer\s+posting|new\s+delivery|Lieferschein)\s+(\d{7,12})',
    re.IGNORECASE,
)
_SAP_OSS_RE      = re.compile(r'OSS\s*Note\s*([\d/]+)', re.IGNORECASE)


# Function: _extract_sap_entities
def _extract_sap_entities(text: str) -> dict:
    """
    Extract SAP-specific identifiers from document text and return them as a
    dict for merging into vector metadata. The retrieval layer uses these
    fields for exact-ID lookups (Layer 0) before falling back to embeddings.
    """
    entities: dict = {}

    incs = list(dict.fromkeys(m.upper() for m in _SAP_INC_RE.findall(text)))
    if incs:
        entities["incident_number"]  = incs[0]
        entities["incident_numbers"] = ",".join(incs)

    solman = _SAP_SOLMAN_RE.findall(text)
    if solman:
        entities["solman_id"] = solman[0]

    deliveries = list(dict.fromkeys(_SAP_DELIVERY_RE.findall(text)))
    if deliveries:
        entities["delivery_numbers"] = ",".join(deliveries)

    oss = list(dict.fromkeys(_SAP_OSS_RE.findall(text)))
    if oss:
        entities["oss_notes"] = ",".join(oss)

    return entities

# ── Regex patterns for incident record detection ──────────────────────────
# Paragraph starts a new incident record (INC number or Analysis document header)
_INC_HEADER_RE = re.compile(
    r'^\s*(INC\d+|Incident\s*(No\.?|Number\s*:?|:)?\s*INC\d+)',
    re.IGNORECASE,
)
# Paragraph is an issue / problem / description label
_ISSUE_RE = re.compile(
    r'^\s*(Issue|Problem|Description)\s*[:\-\.]',
    re.IGNORECASE,
)
# Paragraph marks the start of a solution / resolution section
_SOLUTION_RE = re.compile(
    r'^\s*(Solution|Resolution|Answer|Fix|Workaround|Action\s+Taken|Corrective\s+Action)\s*[:\-\.]?',
    re.IGNORECASE,
)
# Other named sections that are NOT incident boundaries and NOT solutions
_SECTION_RE = re.compile(
    r'^\s*(Analysis|Root\s*Cause|Cause|Impact|Environment|Steps|Scenario|Notes?)\s*[:\-\.]',
    re.IGNORECASE,
)


# ── DOCX ─────────────────────────────────────────────────────

# Function: _is_bold_para
def _is_bold_para(para) -> bool:
    """Return True when ALL non-empty text runs in the paragraph are bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


# Function: _classify_docx_paragraphs
def _classify_docx_paragraphs(doc) -> List[tuple]:
    classified: List[tuple] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.lower().startswith("heading"):
            classified.append(("record_start", text))
            continue
        is_bold = _is_bold_para(para)
        if _INC_HEADER_RE.match(text):
            classified.append(("record_start", text))
        elif _ISSUE_RE.match(text):
            classified.append(("issue", text))
        elif _SOLUTION_RE.match(text):
            classified.append(("solution", text))
        elif _SECTION_RE.match(text) or (is_bold and len(text) < 100):
            classified.append(("section", text))
        else:
            classified.append(("body", text))
    return classified


# Function: _group_docx_records
def _group_docx_records(classified: List[tuple]) -> List[List[tuple]]:
    has_record_starts = any(t == "record_start" for t, _ in classified)
    has_issue_markers = any(t == "issue" for t, _ in classified)
    if has_record_starts:
        boundary_tag = "record_start"
    elif has_issue_markers:
        boundary_tag = "issue"
    else:
        boundary_tag = None
    records: List[List[tuple]] = []
    current: List[tuple] = []
    if boundary_tag:
        for tag, text in classified:
            if tag == boundary_tag and current:
                records.append(current)
                current = [(tag, text)]
            else:
                current.append((tag, text))
        if current:
            records.append(current)
    else:
        records = [classified]
    return records


# Function: _load_docx
def _load_docx(path: Path) -> List[Document]:
    """
    Pattern-based + regex incident parser for DOCX files.

    These documents use bold plain text / paragraph content (NOT Word heading
    styles) to label sections.  This function detects boundaries using:
      1. Regex matching against paragraph text  (INC numbers, "Issue:", "Solution:", …)
      2. Bold paragraph detection               (all-bold short paragraph → section label)
      3. Word heading styles as a tertiary hint

    Each complete incident record (e.g. INC header + Issue + Solution) is
    returned as ONE atomic Document so that the issue description and its
    solution always appear together and are never split by the text splitter.
    """
    doc = DocxDocument(str(path))
    classified = _classify_docx_paragraphs(doc)
    if not classified:
        return []
    records = _group_docx_records(classified)
    result: List[Document] = []
    for group in records:
        content = "\n".join(t for _, t in group).strip()
        if not content:
            continue
        has_solution = any(tag == "solution" for tag, _ in group)
        sap = _extract_sap_entities(content)
        result.append(Document(
            page_content=content,
            metadata={
                "source": path.name,
                "type": "docx",
                "atomic": True,           # Tell the splitter: do NOT cut this
                "has_solution": has_solution,
                **sap,
            },
        ))
    if not result:
        all_text = "\n".join(t for _, t in classified)
        return [Document(page_content=all_text, metadata={"source": path.name, "type": "docx"})]
    logger.info(
        "Parsed %d incident records from '%s' (%d with explicit solution label)",
        len(result), path.name, sum(1 for d in result if d.metadata.get("has_solution")),
    )
    return result


# ── XLSX / CSV ────────────────────────────────────────────────

# Column-name aliases used to detect each semantic role in Excel files
_ISSUE_COLS = {"short description", "short_description", "description", "issue", "problem", "summary", "title"}
_RESOLUTION_COLS = {"resolution", "solution", "answer", "fix", "close notes", "close_notes", "resolution notes"}
_WORKNOTE_COLS = {"work notes", "work_notes", "notes", "comments", "activity", "journal"}
_INC_NUM_COLS = {"number", "incident", "inc", "incident number", "ticket", "id"}
_STATE_COLS = {"state", "status", "incident state"}

# Pattern for ServiceNow-style timestamped work-note entries
_WN_ENTRY_RE = re.compile(
    r'\d{2}\.\d{2}\.\d{4}\s+\d{2}:\d{2}:\d{2}\s*-\s*.+?\(Work notes\)\n',
    re.IGNORECASE,
)
# Patterns for lines that are pure noise (no technical content)
_WN_NOISE_RE = re.compile(
    r'^('
    r'Attachment\s+(added|removed):'
    r'|Dear\s+(Team|FLS|All)[,\s]'
    r'|(Hi|Hello|Dear)\s+\w+\s*[,\n]'  # "Hi Patrick,"
    r'|Gentle\s+Reminder'
    r'|Please\s+find\s+attached'
    r'|Thanks?\s+(and\s+Regards?|Regards?)?\.?\s*$'
    r'|\s*$'
    r')',
    re.IGNORECASE | re.MULTILINE,
)


# Function: _col
def _col(df_cols: list, aliases: set) -> str | None:
    """Return the first DataFrame column name matching any alias (case-insensitive)."""
    for c in df_cols:
        if str(c).strip().lower() in aliases:
            return c
    return None


# Function: _parse_work_notes
def _parse_work_notes(raw: str) -> tuple[str, str]:
    """
    Split a ServiceNow Work notes field into (full_notes, extracted_resolution).

    Work notes in ServiceNow exports are NEWEST-FIRST.  The resolution is taken
    from the MOST RECENT entry that contains substantive technical analysis
    (more than 80 characters after stripping greeting/noise lines).

    Returns (clean_notes, resolution_text).
    """
    if not raw or not raw.strip():
        return "", ""

    # Split into timestamped entries; keep the header as part of each entry
    raw = raw.strip()
    entries = _WN_ENTRY_RE.split(raw)
    headers = _WN_ENTRY_RE.findall(raw)

    # Re-associate header + body.
    # re.split() produces: [pre-match, body1, body2, ...]
    # re.findall() produces: [header1, header2, ...]
    # So entries[i] (for i>=1) matches headers[i-1].
    combined: list[str] = []
    for i, body in enumerate(entries):
        body = body.strip()
        if not body:
            continue
        # entries[0] is text before the first timestamp (usually empty)
        header = headers[i - 1].strip() if i > 0 and (i - 1) < len(headers) else ""
        combined.append(f"{header}\n{body}".strip() if header else body)

    if not combined:
        return raw.strip(), ""

    # Work notes are NEWEST-FIRST in ServiceNow exports.
    # Iterate forward (combined[0] = most recent entry) and pick the first
    # entry with substantive technical content as the resolution.
    resolution = ""
    for entry in combined:
        # Split off the header line to evaluate only the body
        lines = entry.split("\n", 1)
        body = lines[1].strip() if len(lines) > 1 else entry.strip()
        # Remove known noise lines from the body text for evaluation
        clean_lines = [
            ln for ln in body.splitlines()
            if ln.strip() and not _WN_NOISE_RE.match(ln.strip())
        ]
        clean_body = "\n".join(clean_lines).strip()
        if len(clean_body) > 80:          # at least 80 chars of real content
            resolution = clean_body
            break

    clean_notes = "\n\n".join(combined)
    return clean_notes, resolution


# Function: _extract_row_field_values
def _extract_row_field_values(row: pd.Series, columns: list) -> dict:
    inc_col = _col(columns, _INC_NUM_COLS)
    issue_col = _col(columns, _ISSUE_COLS)
    res_col = _col(columns, _RESOLUTION_COLS)
    wn_col = _col(columns, _WORKNOTE_COLS)
    state_col = _col(columns, _STATE_COLS)
    assigned_col = _col(columns, {"assigned to", "assignedto", "assigned_to", "owner"})
    return {
        "inc_col": inc_col,
        "issue_col": issue_col,
        "res_col": res_col,
        "wn_col": wn_col,
        "state_col": state_col,
        "assigned_col": assigned_col,
        "inc_num": str(row[inc_col]).strip() if inc_col and pd.notna(row.get(inc_col, "")) else "",
        "issue": str(row[issue_col]).strip() if issue_col and pd.notna(row.get(issue_col, "")) else "",
        "resolution_col_val": str(row[res_col]).strip() if res_col and pd.notna(row.get(res_col, "")) else "",
        "work_notes_raw": str(row[wn_col]).strip() if wn_col and pd.notna(row.get(wn_col, "")) else "",
        "state": str(row[state_col]).strip() if state_col and pd.notna(row.get(state_col, "")) else "",
        "assigned": str(row[assigned_col]).strip() if assigned_col and pd.notna(row.get(assigned_col, "")) else "",
    }


# Function: _row_to_document
def _row_to_document(row: pd.Series, columns: list, source: str) -> Document | None:
    """
    Convert one Excel row into an atomic incident Document.
    Structure: INC header + Issue: + State + Resolution: (if available) + Work Notes.
    Returns None if the row has no meaningful content.
    """
    f = _extract_row_field_values(row, columns)
    inc_col = f["inc_col"]
    issue_col = f["issue_col"]
    res_col = f["res_col"]
    wn_col = f["wn_col"]
    state_col = f["state_col"]
    assigned_col = f["assigned_col"]
    inc_num = f["inc_num"]
    issue = f["issue"]
    resolution_col_val = f["resolution_col_val"]
    work_notes_raw = f["work_notes_raw"]
    state = f["state"]
    assigned = f["assigned"]

    if not issue and not inc_num:
        return None

    # Parse work notes to extract a clean resolution
    clean_notes, wn_resolution = _parse_work_notes(work_notes_raw)
    resolution = resolution_col_val or wn_resolution

    parts: list[str] = []

    # Header line
    header = f"{inc_num} – {issue}" if inc_num and issue else (inc_num or issue)
    parts.append(header)

    # Structured fields
    if issue:
        parts.append(f"Issue: {issue}")
    if state:
        parts.append(f"State: {state}")
    if assigned:
        parts.append(f"Assigned to: {assigned}")

    # Resolution (labelled so the regex extractor can find it)
    has_solution = False
    if resolution:
        parts.append(f"\nResolution: {resolution}")
        has_solution = True

    # Full work notes (for additional context the LLM can scan)
    if clean_notes:
        parts.append(f"\nWork Notes:\n{clean_notes}")

    # Fallback: include other non-empty columns not already captured
    skip_cols = {inc_col, issue_col, res_col, wn_col, state_col, assigned_col}
    extras = []
    for col in columns:
        if col in skip_cols:
            continue
        val = row.get(col, "")
        if pd.notna(val) and str(val).strip():
            extras.append(f"{col}: {val}")
    if extras:
        parts.append("\n" + " | ".join(extras))

    content = "\n".join(parts).strip()
    if not content:
        return None

    sap = _extract_sap_entities(content)
    sap.setdefault("incident_number", inc_num)   # prefer explicit column value
    return Document(
        page_content=content,
        metadata={
            "source": source,
            "type": "xlsx",
            "atomic": True,
            "has_solution": has_solution,
            **sap,
        },
    )


# Function: _load_xlsx
def _load_xlsx(path: Path) -> List[Document]:
    """
    Load an Excel file row-by-row, producing one atomic Document per incident.
    Each document keeps Issue + Resolution together so the splitter never
    separates them.  Timestamped work notes are parsed to extract the last
    substantive resolution entry.
    """
    docs: List[Document] = []
    try:
        with pd.ExcelFile(str(path)) as xl:
            for sheet in xl.sheet_names:
                df = xl.parse(sheet).fillna("")
                columns = list(df.columns)
                row_docs = 0
                for _, row in df.iterrows():
                    doc = _row_to_document(row, columns, path.name)
                    if doc:
                        docs.append(doc)
                        row_docs += 1
                logger.info(
                    "Parsed %d incident rows from '%s' sheet '%s' (%d with resolution)",
                    row_docs, path.name, sheet,
                    sum(1 for d in docs[-row_docs:] if d.metadata.get("has_solution")),
                )
    except Exception as exc:
        logger.warning("Failed to load %s: %s", path.name, exc)
    return docs


# Function: _load_csv
def _load_csv(path: Path) -> List[Document]:
    """
    Load a CSV file using the same row-by-row incident parser as XLSX.

    If the CSV contains recognisable incident columns (number, description,
    resolution, work notes), each row becomes one atomic Document with the
    same Issue + Resolution structure as DOCX/XLSX records.

    Falls back to a plain columnar text dump if no incident columns are found
    (e.g., a reference data CSV with no INC structure).
    """
    try:
        df = pd.read_csv(str(path)).fillna("")
        columns = list(df.columns)

        # Check whether this CSV has incident-style columns
        has_issue_col = _col(columns, _ISSUE_COLS) is not None
        has_inc_col   = _col(columns, _INC_NUM_COLS) is not None

        if has_issue_col or has_inc_col:
            # Structured incident CSV — use the same row→Document logic as XLSX
            docs: List[Document] = []
            row_docs = 0
            for _, row in df.iterrows():
                doc = _row_to_document(row, columns, path.name)
                if doc:
                    doc.metadata["type"] = "csv"
                    docs.append(doc)
                    row_docs += 1
            logger.info(
                "Parsed %d incident rows from CSV '%s' (%d with resolution)",
                row_docs, path.name,
                sum(1 for d in docs if d.metadata.get("has_solution")),
            )
            return docs

        # Generic CSV — produce a human-readable columnar text block
        lines: List[str] = []
        for _, row in df.iterrows():
            parts = [
                f"{col}: {val}"
                for col, val in row.items()
                if str(val).strip() and str(val).strip().lower() not in ("nan", "none", "")
            ]
            if parts:
                lines.append(" | ".join(parts))
        text = "\n".join(lines).strip()
        if text:
            return [Document(
                page_content=text,
                metadata={"source": path.name, "type": "csv"},
            )]
    except Exception as exc:
        logger.warning("Failed to load %s: %s", path.name, exc)
    return []


# ── MD (structure-aware incident parser) ─────────────────────

_MD_RESOLUTION_HEADING_RE = re.compile(
    r'^(?:Resolution|Solution|Fix|Workaround|Action\s+Taken|Corrective\s+Action)$',
    re.IGNORECASE,
)
_MD_ISSUE_HEADING_RE = re.compile(
    r'^(?:Summary|Description|Issue|Problem|Background)$',
    re.IGNORECASE,
)
_MD_ROOTCAUSE_HEADING_RE = re.compile(
    r'^(?:Root\s*Cause|Cause|Analysis)$',
    re.IGNORECASE,
)


# Function: _build_md_document
def _build_md_document(block: str, path: Path) -> "Document | None":
    block = block.strip()
    if not block:
        return None
    h1_match = re.match(r'^#\s+(.+)', block)
    title = h1_match.group(1).strip() if h1_match else ""
    parts: list[str] = []
    if title:
        parts.append(title)
    section_pieces = re.split(r'(?m)^(?=##\s)', block)
    has_solution = False
    for piece in section_pieces:
        h2_match = re.match(r'^##\s+(.+)\n', piece + "\n")
        if not h2_match:
            leftover = re.sub(r'^#[^\n]*\n?', '', piece).strip()
            if leftover:
                parts.append(leftover)
            continue
        heading = h2_match.group(1).strip()
        body = piece[h2_match.end():].strip()
        body = re.sub(r'\*\*([^*]+)\*\*', r'\1', body)
        body = re.sub(r'(?m)^-{3,}\s*$', '', body).strip()
        if not body:
            continue
        if _MD_RESOLUTION_HEADING_RE.match(heading):
            parts.append(f"\nResolution: {body}")
            has_solution = True
        elif _MD_ISSUE_HEADING_RE.match(heading):
            parts.append(f"Issue: {body}")
        elif _MD_ROOTCAUSE_HEADING_RE.match(heading):
            parts.append(f"Root Cause: {body}")
        elif re.match(r'^Keywords?$', heading, re.IGNORECASE):
            parts.append(f"Keywords: {body}")
        else:
            parts.append(f"{heading}:\n{body}")
    text = "\n".join(parts).strip()
    if not text:
        return None
    sap = _extract_sap_entities(text)
    return Document(
        page_content=text,
        metadata={
            "source": path.name,
            "type": "md",
            "atomic": True,
            "has_solution": has_solution,
            **sap,
        },
    )


# Function: _load_md
def _load_md(path: Path) -> List[Document]:
    """
    Structure-aware parser for Markdown incident files.

    Converts each H2 section into a flat labeled-text block so that the
    downstream regex extractors (Resolution:, Issue:, Root Cause:) identify
    sections without needing to understand Markdown syntax.

    Pass 1 — split the file on H1 boundaries (each H1 = one incident).
    Pass 2 — within each H1 block, split on H2 section headings and map
              the heading to the labelled flat-text format the pipeline uses.
    Pass 3 — extract SAP entities (INC numbers, SolManID, delivery numbers)
              and store them in metadata for exact-ID lookup (Layer 0).

    Every resulting Document is marked atomic=True so the splitter does NOT
    cut it further — issue + resolution always stay in the same chunk.
    """
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as exc:
        logger.warning("Failed to read %s: %s", path.name, exc)
        return []
    if not raw:
        return []
    h1_blocks = re.split(r'(?m)^(?=#\s)', raw)
    results: List[Document] = []
    for block in h1_blocks:
        doc = _build_md_document(block, path)
        if doc:
            results.append(doc)
    if results:
        logger.info(
            "Parsed %d incident record(s) from '%s' (%d with resolution)",
            len(results), path.name,
            sum(1 for d in results if d.metadata.get("has_solution")),
        )
    return results


# ── TXT ──────────────────────────────────────────────────────

# Function: _load_txt
def _load_txt(path: Path) -> List[Document]:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            return [Document(page_content=text, metadata={"source": path.name, "type": "txt"})]
    except Exception as exc:
        logger.warning("Failed to load %s: %s", path.name, exc)
    return []


# ── PDF ──────────────────────────────────────────────────────

# Function: _load_pdf
def _load_pdf(path: Path) -> List[Document]:
    try:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(str(path))
        pages = loader.load()
        for p in pages:
            p.metadata["source"] = path.name
            p.metadata["type"] = "pdf"
        return pages
    except Exception as exc:
        logger.warning("PDF load failed for %s: %s", path.name, exc)
    return []


# ── Image / OCR ───────────────────────────────────────────────

# Function: _ocr_image_bytes
def _ocr_image_bytes(data: bytes, filename: str = "upload") -> str:
    """Run OCR on raw image bytes; tries pytesseract first, then easyocr."""
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    # Try tesseract first
    try:
        import pytesseract
        return pytesseract.image_to_string(img)
    except Exception:
        pass
    # Fallback: easyocr
    try:
        import easyocr
        import numpy as np
        reader = easyocr.Reader(["en"], gpu=False)
        result = reader.readtext(np.array(img), detail=0)
        return "\n".join(result)
    except Exception as exc:
        logger.warning("OCR failed for %s: %s", filename, exc)
    return ""


# Function: load_image_file
def load_image_file(path: Path) -> List[Document]:
    data = path.read_bytes()
    text = _ocr_image_bytes(data, path.name)
    if text.strip():
        return [Document(page_content=text, metadata={"source": path.name, "type": "image_ocr"})]
    return []


# Function: load_image_bytes
def load_image_bytes(data: bytes, filename: str) -> List[Document]:
    text = _ocr_image_bytes(data, filename)
    if text.strip():
        return [Document(page_content=text, metadata={"source": filename, "type": "image_ocr"})]
    return []


# ── Public API ────────────────────────────────────────────────

# Function: load_file
def load_file(path: Path) -> List[Document]:
    suffix = path.suffix.lower()
    dispatch = {
        ".docx": _load_docx,
        ".xlsx": _load_xlsx,
        ".xls":  _load_xlsx,
        ".csv":  _load_csv,
        ".txt":  _load_txt,
        ".md":   _load_md,      # structure-aware incident parser
        ".pdf":  _load_pdf,
        ".png":  load_image_file,
        ".jpg":  load_image_file,
        ".jpeg": load_image_file,
    }
    loader = dispatch.get(suffix)
    if loader is None:
        logger.info("Unsupported file type skipped: %s", path.name)
        return []
    return loader(path)


# Function: load_directory
def load_directory(directory: str) -> List[Document]:
    dir_path = Path(directory)
    if not dir_path.exists():
        logger.warning("Data directory not found: %s", directory)
        return []
    docs: List[Document] = []
    for fp in sorted(dir_path.iterdir()):
        if fp.is_file() and not fp.name.startswith("~$"):
            loaded = load_file(fp)
            docs.extend(loaded)
            logger.info("Loaded %d chunk(s) from %s", len(loaded), fp.name)
    logger.info("Total documents loaded: %d", len(docs))
    return docs
