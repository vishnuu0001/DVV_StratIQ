# ---------------------------------------------------------------------------
# Author: Vishnuu A
# Scope: Parses the two reference contract exhibits (Harmonization Services.docx,
# Date: 2025-10-27
# ---------------------------------------------------------------------------
"""
Parses the two reference contract exhibits (Harmonization Services.docx,
Modernization Services.docx) once, lazily, and exposes small lookup functions
used to ground LLM prompts and the stage-breakdown drill-down endpoint in the
actual contractual methodology rather than inventing lifecycle language.

Parsing walks the document body in order, tracking the current Heading 3 as the
active lifecycle stage, and reads only tables whose first cell is "Sec." (the
cover-page metadata table in both docs fails this check and is skipped).
"""
from __future__ import annotations

import functools
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from config import HARMONIZATION_DOCX_PATH, MODERNIZATION_DOCX_PATH

# Scope-column header abbreviation (as printed in the docx table) -> canonical name
# used everywhere else in the app (VALID_MIGRATION_TYPES / VALID_SERVICE_MODULES).
_HARMONIZATION_SCOPE_COLS = {
    "Data Migr.": "Data Migration Only",
    "Func. Migr.": "Functional Migration Only",
    "Full Con.": "Full Consolidation",
}
_MODERNIZATION_SCOPE_COLS = {
    "Re-host": "Rehost",
    "Re-plat.": "Replatform",
    "Re-fact.": "Refactor",
    "Re-plac.": "Replace",
}

# Verbatim from Modernization Services.docx, "Purpose and Context".
GATE_TEXT = (
    "Modernization is treated as an exception case, applicable only where harmonization onto "
    "strategic target platforms is not feasible and where at least one approved trigger category "
    "is satisfied (IT security risk, legal/compliance requirement, technology end-of-life, or "
    "structural necessity)."
)

TYPE_DEFINITIONS = {
    "Data Migration Only": (
        "Transfer of data from source systems to target platforms without migration of application "
        "functionality or business processes; source application is retired or its functionality is "
        "absorbed by existing target platform capabilities."
    ),
    "Functional Migration Only": (
        "Migration of business processes, configurations, and users from source applications to target "
        "platforms without significant data transformation; applies when data structures are compatible "
        "or data remains in place while process execution shifts to the target platform."
    ),
    "Full Consolidation": (
        "Comprehensive migration combining both data and functional migration, including process "
        "harmonization and integration redesign; applies to complex scenarios, often multiple source "
        "systems consolidating into a single target platform."
    ),
    "Rehost": (
        'Migration of applications to new infrastructure without code changes ("lift and shift"), moving '
        "applications between environments while maintaining the existing application architecture and codebase."
    ),
    "Replatform": (
        "Upgrade of technology stack (frameworks, operating systems, databases) while maintaining existing "
        "application architecture; extends application lifespan by addressing technical debt and restoring vendor support."
    ),
    "Refactor": (
        "Restructuring and optimization of application code and architecture for cloud-native patterns; "
        "includes containerization, API modernization, DevOps enablement, up to microservices decomposition."
    ),
    "Replace": (
        "Substitution of the existing application with an alternative solution (SaaS, COTS product, or "
        "different platform); includes vendor selection, data/user migration, training, and legacy decommissioning."
    ),
}

_HARMONIZATION_STAGE_ORDER = [
    "Assessment Services and Business Case Validation",
    "Architecture and Roadmap Services",
    "Migration and Consolidation Services",
    "Change Management Services (for selected applications)",
    "Post-Harmonization Support Services",
    "Decommissioning Services",
]
_MODERNIZATION_STAGE_ORDER = [
    "Assessment Services and Business Case Validation",
    "Architecture and Roadmap Services",
    "Infrastructure Migration Services",
    "Development Services",
    "Vendor Selection and Solution Configuration Services",
    "Data Migration Services",
    "Testing and Quality Assurance Services",
    "Change Management Services (for selected applications)",
    "Deployment and Transition Services",
    "Decommissioning Services",
]


# Function: _iter_block_items
def _iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


# Function: _handle_stage_paragraph
def _handle_stage_paragraph(item: Paragraph, stages: dict[str, dict[str, Any]], state: dict[str, Any]) -> None:
    style = item.style.name if item.style else ""
    text = item.text.strip()
    if style == "Heading 2" and text == "Activities and Detailed Responsibilities":
        state["in_activities_section"] = True
        return
    if style == "Heading 1" and text == "Service Level Agreement":
        state["in_activities_section"] = False
        return
    if style == "Heading 3" and state["in_activities_section"]:
        state["current_stage"] = text
        stages[text] = {"description": "", "requirements": []}
        return
    current_stage = state["current_stage"]
    if current_stage and text and not stages[current_stage]["requirements"]:
        # Narrative paragraph appearing between the Heading 3 and its table, if any.
        stages[current_stage]["description"] = (stages[current_stage]["description"] + " " + text).strip()


# Function: _requirement_applies_to
def _requirement_applies_to(cells: list[str], scope_header: list[str], scope_cols: dict[str, str], n_scope: int) -> list[str]:
    return [
        scope_cols[scope_header[i]]
        for i in range(1, 1 + n_scope)
        if i < len(cells) and i < len(scope_header)
        and cells[i].strip().upper() == "X" and scope_header[i] in scope_cols
    ]


# Function: _append_requirement_row
def _append_requirement_row(
    row, stage_entry: dict[str, Any], scope_header: list[str], scope_cols: dict[str, str],
    n_scope: int, content_col: int,
) -> None:
    cells = [c.text.strip() for c in row.cells]
    if len(cells) <= content_col or not cells[0]:
        return
    applies_to = _requirement_applies_to(cells, scope_header, scope_cols, n_scope)
    stage_entry["requirements"].append({
        "sec": cells[0],
        "content": cells[content_col],
        "applies_to": applies_to,
    })


# Function: _handle_stage_table
def _handle_stage_table(
    item: Table, stages: dict[str, dict[str, Any]], scope_cols: dict[str, str], state: dict[str, Any]
) -> None:
    current_stage = state["current_stage"]
    if not current_stage or current_stage not in stages:
        return
    rows = item.rows
    if not rows:
        return
    header = [c.text.strip() for c in rows[0].cells]
    if not header or header[0] != "Sec.":
        return  # cover-page metadata table, not a requirements table
    if len(rows) < 2:
        return
    scope_header = [c.text.strip() for c in rows[1].cells]
    n_scope = len(scope_cols)
    content_col = 1 + n_scope
    for row in rows[2:]:
        _append_requirement_row(row, stages[current_stage], scope_header, scope_cols, n_scope, content_col)


# Function: _parse_doc
def _parse_doc(path: str, scope_cols: dict[str, str]) -> dict[str, dict[str, Any]]:
    """Returns {stage_name: {"description": str, "requirements": [{"sec","content","applies_to"}]}}."""
    doc = Document(path)
    stages: dict[str, dict[str, Any]] = {}
    state: dict[str, Any] = {"current_stage": None, "in_activities_section": False}

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            _handle_stage_paragraph(item, stages, state)
        elif isinstance(item, Table):
            _handle_stage_table(item, stages, scope_cols, state)
    return stages


# Function: _harmonization_stages
@functools.lru_cache(maxsize=1)
def _harmonization_stages() -> dict[str, dict[str, Any]]:
    return _parse_doc(HARMONIZATION_DOCX_PATH, _HARMONIZATION_SCOPE_COLS)


# Function: _modernization_stages
@functools.lru_cache(maxsize=1)
def _modernization_stages() -> dict[str, dict[str, Any]]:
    return _parse_doc(MODERNIZATION_DOCX_PATH, _MODERNIZATION_SCOPE_COLS)


# Function: _stages_for
def _stages_for(disposition: str) -> tuple[dict[str, dict[str, Any]], list[str]]:
    if disposition == "Modernization":
        return _modernization_stages(), _MODERNIZATION_STAGE_ORDER
    return _harmonization_stages(), _HARMONIZATION_STAGE_ORDER


# Function: get_disposition_gate_text
def get_disposition_gate_text() -> str:
    return GATE_TEXT


# Function: get_type_definitions
def get_type_definitions(disposition: str) -> dict[str, str]:
    names = list(_MODERNIZATION_SCOPE_COLS.values()) if disposition == "Modernization" else list(_HARMONIZATION_SCOPE_COLS.values())
    return {name: TYPE_DEFINITIONS[name] for name in names}


# Function: get_stage_context
def get_stage_context(disposition: str, max_rows: int = 5, max_chars: int = 220) -> str:
    """Small, fixed-size Assessment-stage snippet block shared by every app in a
    classification batch — used instead of a per-type slice because at classification
    time the type/module hasn't been decided yet (chicken-and-egg). Rows that apply to
    fewer types are more differentiating and are preferred."""
    try:
        stages, order = _stages_for(disposition)
        stage = stages.get(order[0], {})  # "Assessment Services and Business Case Validation"
        reqs = stage.get("requirements", [])
        reqs_sorted = sorted(reqs, key=lambda r: len(r["applies_to"]) or 99)
        lines = []
        for r in reqs_sorted[:max_rows]:
            tag = f"[{', '.join(r['applies_to'])}] " if r["applies_to"] else ""
            lines.append(f"- {r['sec']}: {tag}{r['content'][:max_chars]}")
        return "\n".join(lines)
    except Exception:
        return ""


# Function: get_classification_reference_block
def get_classification_reference_block(max_chars: int = 800) -> str:
    """Combined, capped Assessment-stage snippet covering both Harmonization migration
    types and Modernization service modules — used as shared context for the per-app
    disposition + migration-type/service-module classification prompt. Fixed/shared
    across every app in a batch because the type/module isn't known until after
    classification (see get_stage_context's chicken-and-egg note)."""
    harmon = get_stage_context("Harmonization", max_rows=3, max_chars=140)
    modern = get_stage_context("Modernization", max_rows=3, max_chars=140)
    lines = (harmon + "\n" + modern).split("\n")
    out_lines: list[str] = []
    total = 0
    for line in lines:
        if total + len(line) + 1 > max_chars:
            break
        out_lines.append(line)
        total += len(line) + 1
    return "\n".join(out_lines)


_DELIVERY_MODEL_STAGES = ("Wave Planning and Prioritization", "Wave Execution")


# Function: get_wave_delivery_context
@functools.lru_cache(maxsize=1)
def get_wave_delivery_context(max_chars: int = 700) -> str:
    """Shared 'Delivery Model (Wave-Based)' narrative — identical structure in both
    docs — used as grounding context for wave_assignment_rationale generation."""
    try:
        doc = Document(HARMONIZATION_DOCX_PATH)
        capture = False
        lines: list[str] = []
        for p in doc.paragraphs:
            style = p.style.name if p.style else ""
            text = p.text.strip()
            if style == "Heading 3" and text in _DELIVERY_MODEL_STAGES:
                capture = True
                continue
            if style in ("Heading 1", "Heading 2"):
                capture = False
            elif style == "Heading 3" and text not in _DELIVERY_MODEL_STAGES:
                capture = False
            if capture and text:
                lines.append(text)
        return " ".join(lines)[:max_chars]
    except Exception:
        return ""


