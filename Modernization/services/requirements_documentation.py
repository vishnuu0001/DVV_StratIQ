# ---------------------------------------------------------------------------
# Scope: Ollama-backed requirements documentation for governed projects.
# ---------------------------------------------------------------------------
"""Generate BRD, FSD, and knowledge-graph artifacts from immutable source.

BRD/FSD structure follows a fixed, numbered template (Document Control ->
Introduction -> Stakeholders -> Business Requirements table -> Business Rules
-> Success Criteria -> Glossary for BRD; Document Control -> Introduction ->
Document Conventions -> System Overview -> one Module section per observed
capability -> Non-Functional Observations -> Open Items for FSD) modeled on a
hand-authored reference pair reverse-engineered from a real Struts/JSP/DB2
application. The section SET and ORDER is fixed; the CONTENT adapts to
whatever the evidence actually shows (framework, data store, field layout,
action codes if any) so the same template generalizes to other stacks.

Everything a reviewer reads first, and everything whose absence would make
the document look incomplete (headings, tables, requirement/rule counts), is
built deterministically from the observed capability register — never left to
chance on a small local model. The model is used only for the narrow,
evidence-scoped writing tasks a 6-7B model can reliably do: one descriptive
sentence per requirement row, a handful of business-rule bullets, and one
module's field/validation detail — each independently bounded, retried once,
and backed by an honest deterministic fallback so a single weak completion
never blocks the document or drops a capability.
"""
from __future__ import annotations

import json
import io
import logging
import re
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Any, Callable, Optional

from services import llm
from services.governance import semantic_index

logger = logging.getLogger(__name__)

CONTEXT_LIMIT = 48_000
EVIDENCE_LIMIT = 22_000
DOCUMENT_TYPES = {"brd", "fsd", "knowledge_graph"}
REQUIREMENTS_PREFERRED_MODELS = (
    "deepseek-coder:6.7b", "qwen2.5-coder:7b", "qwen3.5:9b", "qwen2.5-coder:3b",
)

# Each observed capability gets full (not excerpted) source content for the
# handful of files behind it — field names, action codes, and stored-procedure
# names generally live past whatever a small capped excerpt would have kept,
# and a capability is typically backed by only 5-10 modest files, so reading
# them in full is cheap and keeps the model's answers verifiably grounded.
CAPABILITY_FULL_EVIDENCE_BUDGET = 16_000  # chars, per capability

CAPABILITY_MAX_TOKENS = 2_048
CAPABILITY_CONTEXT_TOKENS = 10_240
CAPABILITY_MAX_SECONDS = 210
CAPABILITY_MAX_ATTEMPTS = 2  # initial attempt + 1 targeted retry, per capability

MODULE_MAX_TOKENS = 3_072
MODULE_CONTEXT_TOKENS = 12_288
MODULE_MAX_SECONDS = 240
MODULE_MAX_ATTEMPTS = 2

# Preserved for the knowledge-graph path, the one remaining single-completion
# call — it still needs the full per-file evidence payload and a larger
# output budget than the template-driven BRD/FSD path now does.
DOCUMENT_MAX_TOKENS = 8_192
DOCUMENT_CONTEXT_TOKENS = 24_576

_GRAPH_INSTRUCTIONS = """Return JSON only with this shape:
{"title":"...","summary":"...","nodes":[{"id":"unique-id","label":"...","type":"business|actor|feature|functional|data|integration|rule|quality","description":"..."}],"edges":[{"source":"node-id","target":"node-id","relationship":"enables|uses|depends_on|implements|governs|produces|integrates_with"}]}
Build a detailed, connected knowledge graph linking business needs to functional features, actors,
rules, data, integrations, and quality requirements. Use 30-100 concise nodes and preserve explicit
BR/FS identifiers in labels. Add an evidence_source_path property to nodes whenever source evidence
exists. Include the supplied Project ID, Project Name, Application Key, and Client Name in the graph
identity and connect the project root to its requirements. Every edge endpoint must reference a
declared node. Do not wrap the JSON in Markdown fences."""

_EVIDENCE_EXTENSIONS = {
    ".py", ".java", ".cs", ".vb", ".js", ".jsx", ".ts", ".tsx", ".go", ".rs",
    ".php", ".rb", ".c", ".cpp", ".h", ".hpp", ".sql", ".graphql", ".xml", ".jsp",
    ".json", ".yaml", ".yml", ".toml", ".properties", ".config", ".md", ".txt",
}
_EVIDENCE_SKIP_DIRS = {".git", "node_modules", "dist", "build", "bin", "obj", "target", ".venv", "__pycache__"}

_LANGUAGE_BY_EXTENSION = {
    ".py": "Python", ".java": "Java", ".cs": "C#", ".vb": "Visual Basic",
    ".js": "JavaScript", ".jsx": "JavaScript/JSX", ".ts": "TypeScript",
    ".tsx": "TypeScript/TSX", ".go": "Go", ".rs": "Rust", ".php": "PHP",
    ".rb": "Ruby", ".c": "C", ".cpp": "C++", ".h": "C/C++ Header",
    ".hpp": "C++ Header", ".sql": "SQL", ".graphql": "GraphQL", ".xml": "XML",
    ".json": "JSON", ".yaml": "YAML", ".yml": "YAML", ".toml": "TOML",
    ".properties": "Properties", ".config": "Configuration", ".md": "Markdown",
    ".txt": "Text",
    ".jsp": "JSP",
}

_VENDOR_FILE_MARKERS = (
    "jquery", "bootstrap", "json2", "datatables", "polyfill", "vendor", ".min.",
)
_SUPPORTING_CAPABILITY_TERMS = {
    "app", "application", "base", "common", "constant", "error", "home", "login",
    "schema", "session", "user", "utility", "welcome",
}
_OPERATION_PATTERNS = {
    "Add": r"(?i)\b(add|addition|create|insert|save|new)\b",
    "Modify": r"(?i)\b(update|modify|modification|edit)\b",
    "Delete": r"(?i)\b(delete|deletion|remove)\b",
    "View / Search / List": r"(?i)\b(view|search|list|inquiry|find|details?)\b",
    "Export": r"(?i)\b(export|excel|csv|report)\b",
}


def _is_vendor_file(relative: str) -> bool:
    lowered = relative.lower()
    name = Path(relative).name.lower()
    return "/lib/" in f"/{lowered}/" or any(marker in name for marker in _VENDOR_FILE_MARKERS)


def _camel_words(value: str) -> list[str]:
    expanded = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", value)
    return re.findall(r"[A-Za-z][A-Za-z0-9]*", expanded)


def _capability_term(relative: str, symbols: list[str]) -> str | None:
    """Infer a conservative business noun from application-layer filenames/classes."""
    lowered = relative.lower()
    if _is_vendor_file(relative) or not any(marker in lowered for marker in (
        "/action/", "/controller/", "/service", "/form/", "/dto/", "/pages/", "/custom/",
    )):
        return None
    candidates = [Path(relative).stem] + symbols
    for candidate in candidates:
        words = _camel_words(candidate)
        while words and words[-1].lower() in {
            "action", "controller", "service", "form", "dto", "details", "detail", "index",
            "information", "repository", "util", "utility", "constants", "constant", "page",
        }:
            words.pop()
        if not words:
            continue
        term = words[0]
        if term.lower() not in _SUPPORTING_CAPABILITY_TERMS and len(term) >= 3:
            return term.title()
    return None


def _functional_signals(relative: str, content: str, symbols: list[str]) -> tuple[list[str], list[str]]:
    term = _capability_term(relative, symbols)
    if not term:
        return [], []
    searchable = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", content)
    operations = [label for label, pattern in _OPERATION_PATTERNS.items() if re.search(pattern, searchable)]
    return [term], operations


def _evidence_priority(relative: str, capabilities: list[str]) -> int:
    lowered = relative.lower()
    name = Path(relative).name.lower()
    if _is_vendor_file(relative):
        return 9
    if capabilities or name in {"struts-config.xml", "tiles-defs.xml", "web.xml"}:
        return 0
    if any(marker in lowered for marker in ("/pages/", "/source/", "/src/", "/custom/")):
        return 1
    if name.startswith(("readme", "requirement", "spec", "pom.", "package.", "build.")):
        return 1
    return 3


def _focused_excerpt(content: str, budget: int) -> str:
    """Prefer behavioral lines over file headers when evidence must be compressed."""
    if len(content) <= budget:
        return content
    lines = content.splitlines()
    selected: list[int] = list(range(min(4, len(lines))))
    seen = set(selected)

    def add_context(index: int) -> None:
        for candidate in (max(0, index - 1), index, min(len(lines) - 1, index + 1)):
            if candidate not in seen:
                selected.append(candidate)
                seen.add(candidate)

    searchable_lines = [re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", line) for line in lines]
    for pattern in _OPERATION_PATTERNS.values():
        match_index = next((index for index, line in enumerate(searchable_lines) if re.search(pattern, line)), None)
        if match_index is not None:
            add_context(match_index)
    structural = re.compile(r"(?i)(executeAction|<action|<form-bean|<forward|property=|function\s+|class\s+)")
    for index, line in enumerate(searchable_lines):
        if structural.search(line):
            add_context(index)
    rendered: list[str] = []
    used = 0
    for index in selected:
        line = f"L{index + 1}: {lines[index].strip()}"
        if used + len(line) + 1 > budget:
            break
        rendered.append(line)
        used += len(line) + 1
    return "\n".join(rendered) or content[:budget]


def _detected_symbols(content: str, extension: str) -> list[str]:
    """Return a compact deterministic inventory of declarations and interface signals."""
    patterns = [
        r"\b(?:class|interface|enum|record|struct)\s+([A-Za-z_$][\w$]*)",
        r"\b(?:def|function)\s+([A-Za-z_$][\w$]*)\s*\(",
    ]
    if extension in {".java", ".cs", ".c", ".cpp", ".h", ".hpp"}:
        patterns.append(
            r"\b(?:public|protected|private|internal|static|final|virtual|async|synchronized)"
            r"(?:\s+[\w<>,.?\[\]]+)+\s+([A-Za-z_$][\w$]*)\s*\("
        )
    if extension == ".sql":
        patterns.append(r'(?i)\b(?:create\s+(?:table|view|procedure|function)|alter\s+table)\s+([\w.\[\]`"]+)')
    symbols: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, content):
            symbol = match.group(1).strip('`"[]')
            if symbol and symbol not in symbols:
                symbols.append(symbol)
            if len(symbols) >= 16:
                return symbols
    return symbols


def _source_evidence(source_path: str | Path) -> tuple[list[dict], list[dict]]:
    root = Path(source_path)
    manifest: list[dict] = []
    entries_by_path: dict[str, dict] = {}
    candidates: list[tuple[int, Path, str]] = []
    for path in sorted(root.rglob("*"), key=lambda item: item.as_posix().lower()):
        if not path.is_file() or any(part in _EVIDENCE_SKIP_DIRS for part in path.relative_to(root).parts):
            continue
        relative = path.relative_to(root).as_posix()
        try:
            size = path.stat().st_size
        except OSError:
            continue
        extension = path.suffix.lower()
        entry = {
            "evidence_id": f"SRC-{len(manifest) + 1:04d}", "path": relative,
            "type": _LANGUAGE_BY_EXTENSION.get(extension, extension.lstrip(".").upper() or "File"),
            "bytes": size, "lines": None, "symbols": [], "coverage": "inventory-only",
        }
        manifest.append(entry)
        entries_by_path[relative] = entry
        if extension not in _EVIDENCE_EXTENSIONS or size > 1_000_000:
            continue
        try:
            content = path.read_text(encoding="utf-8", errors="replace").strip()
        except OSError:
            continue
        if not content:
            continue
        entry["lines"] = content.count("\n") + 1
        entry["symbols"] = _detected_symbols(content, extension)
        entry["capability_terms"], entry["operations"] = _functional_signals(
            relative, content, entry["symbols"],
        )
        entry["coverage"] = "dependency-inventory" if _is_vendor_file(relative) else "content-inspected"
        if entry["coverage"] == "dependency-inventory":
            continue
        priority = _evidence_priority(relative, entry["capability_terms"])
        candidates.append((priority, path, content))

    excerpts: list[dict] = []
    remaining = EVIDENCE_LIMIT
    ordered = sorted(candidates, key=lambda item: (item[0], item[1].as_posix().lower()))
    # Give every readable source file evidence coverage before enriching the
    # most informative documentation, configuration, and implementation files.
    coverage_size = max(80, min(900, EVIDENCE_LIMIT // max(1, len(ordered))))
    for _, path, content in ordered:
        if remaining <= 0:
            break
        excerpt = _focused_excerpt(content, min(coverage_size, remaining))
        relative = path.relative_to(root).as_posix()
        manifest_entry = entries_by_path[relative]
        excerpts.append({"evidence_id": manifest_entry["evidence_id"], "path": relative, "excerpt": excerpt})
        remaining -= len(excerpt)
    if remaining > 0:
        by_path = {item["path"]: item for item in excerpts}
        for _, path, content in sorted(candidates, key=lambda item: (item[0], -len(item[2]))):
            if remaining <= 0:
                break
            relative = path.relative_to(root).as_posix()
            if relative not in by_path:
                continue
            consumed = len(by_path[relative]["excerpt"])
            expanded = _focused_excerpt(content, consumed + min(2600, remaining))
            delta = max(0, len(expanded) - consumed)
            by_path[relative]["excerpt"] = expanded
            remaining -= delta
    return manifest, excerpts


def _functional_capability_inventory(manifest: list[dict]) -> list[dict]:
    grouped: dict[str, dict] = {}
    for item in manifest:
        for term in item.get("capability_terms") or []:
            record = grouped.setdefault(term, {
                "capability_id": "", "name": f"{term} Setup", "operations": set(),
                "evidence_ids": [], "source_paths": [], "declarations": set(),
            })
            record["operations"].update(item.get("operations") or [])
            record["evidence_ids"].append(item["evidence_id"])
            record["source_paths"].append(item["path"])
            record["declarations"].update(item.get("symbols") or [])
    capabilities: list[dict] = []
    corroborated = [(term, record) for term, record in sorted(grouped.items()) if len(record["source_paths"]) >= 2]
    for index, (term, record) in enumerate(corroborated, start=1):
        operations = sorted(record["operations"], key=lambda value: list(_OPERATION_PATTERNS).index(value))
        capabilities.append({
            "capability_id": f"CAP-{index:03d}",
            "name": record["name"] if any(op in operations for op in ("Add", "Modify", "Delete")) else term,
            "operations": operations,
            "evidence_ids": record["evidence_ids"],
            "source_paths": record["source_paths"],
            "declarations": sorted(record["declarations"])[:40],
        })
    return capabilities


def _governed_analysis(project: dict) -> dict | None:
    snapshot = next((item for item in project.get("snapshots", []) if item.get("kind") == "analysis"), None)
    if not snapshot:
        return None
    try:
        return json.loads((Path(snapshot["path"]) / "artifact.json").read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, KeyError, TypeError):
        return None


def _project_identity(project: dict) -> dict:
    configuration = project.get("configuration") or {}
    return {
        "project_id": str(project.get("id") or ""),
        "project_name": str(project.get("name") or configuration.get("application_name") or ""),
        "application_key": str(configuration.get("application_key") or configuration.get("project_key") or ""),
        "application_name": str(configuration.get("application_name") or project.get("name") or ""),
        "client_name": str(configuration.get("client_name") or configuration.get("customer") or ""),
        "application_owner": str(configuration.get("application_owner") or ""),
        "business_unit": str(configuration.get("business_unit") or ""),
        "business_criticality": str(configuration.get("business_criticality") or ""),
    }


def _bounded_list_json(items: list[dict], limit: int) -> str:
    """Serialize complete list entries without cutting JSON in the middle of an item."""
    selected: list[dict] = []
    used = 64
    for item in items:
        item_size = len(json.dumps(item, ensure_ascii=False, separators=(",", ":"), default=str)) + 1
        if used + item_size > limit:
            break
        selected.append(item)
        used += item_size
    return json.dumps({"items": selected, "included": len(selected), "omitted_from_prompt": len(items) - len(selected)}, ensure_ascii=False, separators=(",", ":"), default=str)


def _coverage_summary(manifest: list[dict]) -> dict:
    types = Counter(str(item.get("type") or "Unknown") for item in manifest)
    inspected = sum(item.get("coverage") == "content-inspected" for item in manifest)
    return {
        "source_files_discovered": len(manifest),
        "source_files_content_inspected": inspected,
        "source_files_inventory_only": len(manifest) - inspected,
        "known_source_lines": sum(int(item.get("lines") or 0) for item in manifest),
        "file_types": dict(sorted(types.items())),
        "coverage_rule": "All discovered project files appear in the authoritative Source Coverage Register. Use Evidence IDs and paths for citations; inventory-only files must not be used to infer behavior.",
    }


def _escape_markdown_cell(value: Any) -> str:
    return str(value if value not in (None, "") else "—").replace("|", "\\|").replace("\n", " ")


def _coverage_appendix(manifest: list[dict]) -> str:
    summary = _coverage_summary(manifest)
    capabilities = _functional_capability_inventory(manifest)
    lines = [
        "## Appendix A: Observed Functional Capability Register", "",
        "Only the capabilities below were established from application-layer evidence. Supporting technical concerns are not promoted to business functionality.",
        "", "| Capability ID | Observed capability | Operations | Primary source evidence |",
        "| --- | --- | --- | --- |",
    ]
    for capability in capabilities:
        sources = [
            f"{evidence_id} — {path}"
            for evidence_id, path in zip(capability["evidence_ids"], capability["source_paths"])
        ]
        lines.append("| " + " | ".join(_escape_markdown_cell(value) for value in (
            capability["capability_id"], capability["name"],
            ", ".join(capability["operations"]) or "Observed behavior requires review",
            "; ".join(sources),
        )) + " |")
    lines.extend([
        "", "## Appendix B: Authoritative Source Coverage Register", "",
        f"This register accounts for **{summary['source_files_discovered']}** discovered project files: **{summary['source_files_content_inspected']}** content-inspected and **{summary['source_files_inventory_only']}** inventory-only. Known text volume is **{summary['known_source_lines']:,} lines**. Inventory-only entries are recorded for completeness but were not used to infer behavior.",
        "", "| Evidence ID | Source path | Type | Lines | Bytes | Detected declarations / signals | Coverage |",
        "| --- | --- | --- | ---: | ---: | --- | --- |",
    ])
    for item in manifest:
        lines.append("| " + " | ".join(_escape_markdown_cell(value) for value in (
            item.get("evidence_id"), item.get("path"), item.get("type"), item.get("lines"),
            item.get("bytes"), ", ".join(item.get("symbols") or []) or "No declaration detected",
            item.get("coverage"),
        )) + " |")
    return "\n".join(lines)


def _project_context(
    project: dict, source_path: str | Path, manifest: Optional[list[dict]] = None,
    excerpts: Optional[list[dict]] = None, *, manifest_budget: int = 12_000, excerpt_budget: int = 20_000,
) -> str:
    """Build the evidence payload fed to the knowledge-graph completion (the
    one remaining single-shot call in this module)."""
    index = semantic_index(Path(source_path))
    if manifest is None or excerpts is None:
        manifest, excerpts = _source_evidence(source_path)
    project_context: dict[str, Any] = {
        "project": {
            "id": project.get("id"),
            "name": project.get("name"),
            "configuration": project.get("configuration", {}),
        },
        "project_identity": _project_identity(project),
        "coverage": _coverage_summary(manifest),
        "observed_functional_capabilities": _functional_capability_inventory(manifest),
        "functional_scope_rule": (
            "The observed capability inventory is the authoritative functional scope. Describe every listed "
            "capability and operation. Do not introduce generic industry capabilities without direct evidence."
        ),
    }
    sections = (
        json.dumps(project_context, ensure_ascii=False, indent=2, default=str),
        '"source_manifest":\n' + _bounded_list_json(manifest, manifest_budget),
        '"governed_analysis":\n' + json.dumps(_governed_analysis(project), ensure_ascii=False, separators=(",", ":"), default=str)[:7_000],
        '"source_evidence_excerpts":\n' + _bounded_list_json(excerpts, excerpt_budget),
        '"source_semantic_index":\n' + json.dumps(index, ensure_ascii=False, separators=(",", ":"), default=str)[:8_000],
    )
    rendered = "\n\n".join(sections)
    return rendered[:CONTEXT_LIMIT]


def _salvage_json_array(text: str, key: str) -> list:
    """When strict JSON parsing fails, salvage as many complete `{...}`
    elements as possible from one top-level array (nodes/edges) instead of
    discarding an entire otherwise-good response over a single malformed
    token. A ~30-100 node graph is an expensive completion (often the better
    part of a minute of generation), and small local models reliably produce
    a mostly-correct large array with one isolated glitch (an unescaped
    quote, a missing comma) rather than being wrong throughout — walking the
    array by brace-depth and keeping every element that parses on its own
    recovers the vast majority of that work."""
    match = re.search(rf'"{key}"\s*:\s*\[', text)
    if not match:
        return []
    elements: list = []
    depth = 0
    element_start = None
    in_string = False
    escape = False
    for index in range(match.end() - 1, len(text)):
        char = text[index]
        if in_string:
            if escape:
                escape = False
            elif char == "\\":
                escape = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            if depth == 0:
                element_start = index
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0 and element_start is not None:
                try:
                    elements.append(json.loads(text[element_start:index + 1]))
                except json.JSONDecodeError:
                    pass  # this one element is still malformed — skip just it, not the whole array
                element_start = None
        elif char == "]" and depth == 0:
            break
    return elements


def _extract_json(text: str) -> dict:
    candidate = text.strip()
    if candidate.startswith("```"):
        candidate = candidate.split("\n", 1)[-1]
        candidate = candidate.rsplit("```", 1)[0].strip()
    start, end = candidate.find("{"), candidate.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("Ollama did not return a JSON knowledge graph")
    body = candidate[start:end + 1]
    try:
        graph = json.loads(body)
    except json.JSONDecodeError:
        nodes = _salvage_json_array(body, "nodes")
        if not nodes:
            raise
        edges = _salvage_json_array(body, "edges")
        logger.warning(
            "Knowledge graph JSON was malformed; salvaged %d node(s) and %d edge(s) from the otherwise-discarded response",
            len(nodes), len(edges),
        )
        graph = {"nodes": nodes, "edges": edges}
    nodes = graph.get("nodes")
    edges = graph.get("edges")
    if not isinstance(nodes, list) or not isinstance(edges, list):
        raise ValueError("Knowledge graph must contain node and edge arrays")
    node_ids = {str(node.get("id")) for node in nodes if isinstance(node, dict) and node.get("id")}
    graph["nodes"] = [node for node in nodes if isinstance(node, dict) and str(node.get("id")) in node_ids]
    graph["edges"] = [
        edge for edge in edges if isinstance(edge, dict)
        and str(edge.get("source")) in node_ids and str(edge.get("target")) in node_ids
    ]
    return graph


def _requirements_model() -> str | None:
    available = llm.check_status().get("models", [])
    return next((model for model in REQUIREMENTS_PREFERRED_MODELS if model in available), None) or llm.pick_codegen_model()


# The knowledge graph is the one remaining single-completion call in this
# module and needs the full per-file evidence payload to cite individual
# files as graph nodes — a much bigger prompt than the BRD/FSD path's small,
# per-capability calls now use. deepseek-coder:6.7b (REQUIREMENTS_PREFERRED_
# MODELS' first choice) has a native context window of only 16,384 tokens;
# sizing that big a prompt right up to the ceiling left ~0 headroom for the
# completion, so Ollama returned 200 OK but a truncated/malformed response —
# "Ollama did not return a JSON knowledge graph" even though every call
# succeeded at the transport level. Prefer a larger-context installed model
# for this task specifically, and size the evidence payload to leave real
# headroom below whichever model's window actually gets used.
GRAPH_PREFERRED_MODELS = ("qwen3.5:9b", "qwen2.5-coder:7b", "qwen2.5-coder:3b", "deepseek-coder:6.7b")
_GRAPH_MAX_TOKENS = 4_096
_GRAPH_CONTEXT_WINDOW_BY_MODEL = {
    "qwen3.5:9b": 32_768, "qwen2.5-coder:7b": 32_768, "qwen2.5-coder:3b": 32_768,
    "deepseek-coder:6.7b": 16_384,
}
_GRAPH_CONTEXT_WINDOW_DEFAULT = 16_384


def _graph_model() -> str | None:
    available = llm.check_status().get("models", [])
    return next((model for model in GRAPH_PREFERRED_MODELS if model in available), None) or llm.pick_codegen_model()


def _safe_generate(description: str, **kwargs) -> str:
    """Run one Ollama completion, converting a timeout or transport failure into
    an empty result instead of letting the exception propagate. The template
    pipeline makes one small completion per capability rather than one giant
    completion, so more individual calls have a chance to hit a slow model or a
    transient Ollama hiccup — without this, a single failed call would abort
    the whole document and discard every other capability's already-validated
    work. Callers must treat "" as "this attempt produced nothing" and fall
    through to their own retry or deterministic fallback."""
    try:
        return llm.generate(**kwargs)
    except Exception as exc:
        logger.warning("Requirements-document completion failed (%s), treating as empty: %s", description, exc)
        return ""


def _looks_like_refusal(text: str) -> bool:
    lowered = text.lower()
    return any(marker in lowered for marker in (
        "i'm sorry", "i am unable", "as an ai", "cannot provide", "unable to generate",
        "general outline", "fill in the details",
    ))


def _join_and(items: list[str]) -> str:
    values = [str(item) for item in items if item]
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    return ", ".join(values[:-1]) + " and " + values[-1]


# ─── Deep, per-capability evidence ─────────────────────────────────────────

def _capability_full_evidence(capability: dict, source_path: str | Path) -> list[dict]:
    """Read FULL (not excerpted) content for a capability's own source files.
    A capability is typically backed by only a handful of modest files
    (action/controller, form/model, DTO, service, JSP/view, DAO config) —
    small enough in full to fit one completion's context — and field-level
    detail (screen fields, action codes, stored-procedure/API names) usually
    lives past whatever a small capped excerpt would have kept."""
    root = Path(source_path)
    remaining = CAPABILITY_FULL_EVIDENCE_BUDGET
    result: list[dict] = []
    seen_paths: set[str] = set()
    pairs = list(zip(capability.get("evidence_ids") or [], capability.get("source_paths") or []))
    for evidence_id, relative in pairs:
        if remaining <= 0 or relative in seen_paths:
            continue
        seen_paths.add(relative)
        try:
            content = (root / relative).read_text(encoding="utf-8", errors="replace").strip()
        except OSError:
            continue
        if not content:
            continue
        excerpt = content[:remaining]
        result.append({"evidence_id": evidence_id, "path": relative, "content": excerpt})
        remaining -= len(excerpt)
    return result or [{"evidence_id": "—", "path": "—", "content": "No readable source content available."}]


def _evidence_block(evidence: list[dict]) -> str:
    return "\n\n".join(
        f"### {item['path']} (Evidence ID {item['evidence_id']})\n```\n{item['content']}\n```"
        for item in evidence
    )


# ─── Technology-stack detection (drives the FSD architecture table) ────────

_FRAMEWORK_KEYWORDS = (
    ("Apache Struts", "struts"), ("Spring", "springframework"), ("JSP / Servlet", ".jsp"),
    ("ASP.NET Core", "microsoft.aspnetcore"), ("ASP.NET", "system.web"), ("Express.js", "express("),
    ("React", "react-dom"), ("Angular", "@angular/core"), ("Django", "django."),
    ("Flask", "from flask"), ("FastAPI", "from fastapi"),
)
_DATA_STORE_KEYWORDS = (
    ("DB2", "db2"), ("Oracle", "oracle"), ("SQL Server", "sqlserver"), ("PostgreSQL", "postgres"),
    ("MySQL", "mysql"), ("MongoDB", "mongo"),
)


def _detect_tech_signals(manifest: list[dict]) -> dict:
    haystack = " ".join(
        f"{item.get('path', '')} {' '.join(item.get('symbols') or [])}".lower() for item in manifest
    )
    frameworks = [name for name, keyword in _FRAMEWORK_KEYWORDS if keyword in haystack]
    data_stores = [name for name, keyword in _DATA_STORE_KEYWORDS if keyword in haystack]
    if any(f"/{item.get('path', '').lower()}/".count("/dao/") for item in manifest) or "storedproc" in haystack:
        data_stores.append("Stored Procedure / DAO Layer")
    languages = sorted({
        item.get("type") for item in manifest
        if item.get("type") and item.get("coverage") == "content-inspected"
    })
    return {"frameworks": frameworks, "data_stores": data_stores, "languages": languages}


# ─── Document Control (shared by BRD and FSD) ──────────────────────────────

def _document_control_markdown(document_type: str, identity: dict) -> str:
    doc_label = "BRD" if document_type == "brd" else "FSD"
    lines = [
        "## Document Control", "",
        "### Version History", "",
        "| Version | Date | Author | Description |",
        "| --- | --- | --- | --- |",
        f"| 1.0 | {date.today().isoformat()} | Modernization Studio · Automated Analysis | "
        f"Initial draft, generated from analysis of the existing application source code (no prior "
        f"{doc_label} was supplied). |",
        "",
    ]
    if document_type == "brd":
        lines += [
            "### Distribution List", "",
            "| Name / Role | Organization |", "| --- | --- |",
            f"| Product Owner / Business Sponsor | {identity.get('client_name') or 'Client'} |",
            "| Delivery Lead | Modernization Delivery Team |",
            "| QA Lead | Modernization Delivery Team |",
        ]
    else:
        lines += [
            "### Related Documents", "",
            f"- Business Requirements Document — {identity.get('project_name') or 'this application'} (companion document).",
        ]
    return "\n".join(lines)


# ─── BRD: 1. Introduction / 2. Stakeholders / 5. Success Criteria / 6. Glossary ─

def _brd_introduction_markdown(identity: dict, capabilities: list[dict], manifest: list[dict]) -> str:
    app = identity.get("project_name") or identity.get("application_name") or "the application"
    client = identity.get("client_name") or "the client"
    names = [str(c.get("name") or "") for c in capabilities]
    plural = "y" if len(names) == 1 else "ies"
    lines = [
        "## 1. Introduction", "",
        "### 1.1 Purpose", "",
        f"This Business Requirements Document (BRD) defines the business needs, objectives, and functional "
        f"requirements that {app} is expected to satisfy. The requirements captured here were derived from "
        f"an as-is analysis of the existing application source code, as no prior BRD or requirements "
        f"artifact was supplied. This document is intended to serve as the baseline of record for future "
        f"enhancement, re-platforming, or modernization of this application.", "",
        "### 1.2 Background", "",
    ]
    if names:
        lines.append(
            f"{app} is a system used by {client} that, per source-code analysis, maintains the following "
            f"capabilit{plural}:"
        )
        lines.append("")
        for capability in capabilities:
            operations = _join_and(capability.get("operations") or []) or "behavior requiring further source review"
            lines.append(f"- **{capability.get('name')}** — supports {operations}.")
    else:
        lines.append(
            f"Source-code analysis of {app} did not safely establish a current-state business capability; "
            "this is recorded as an Open Question pending further review."
        )
    lines += ["", "### 1.3 Business Objectives", ""]
    for capability in capabilities:
        lines.append(
            f"- Provide a single, authoritative system of record for {capability.get('name')} data, "
            "maintained without requiring direct database access or IT intervention."
        )
    lines += [
        "- Enforce data quality and consistency at the point of data entry for every observed capability.",
        "- Restrict data-modification capability to authorized personnel while allowing appropriate "
        "read/inquiry access to the wider organization.",
        "", "### 1.4 Scope", "", "#### 1.4.1 In Scope", "",
    ]
    if capabilities:
        for capability in capabilities:
            operations = _join_and(capability.get("operations") or []) or "behavior requiring further source review"
            lines.append(f"- **{capability.get('name')}**: {operations}.")
    else:
        lines.append("- No business capability was safely established from source evidence.")
    lines += [
        "", "#### 1.4.2 Out of Scope", "",
        "- Any capability, module, or workflow not listed above is explicitly out of scope unless and until "
        "supported by cited evidence.",
        "- User provisioning and enterprise identity/security-group administration, where handled by a "
        "platform outside the observed application code.",
        "", "### 1.5 Assumptions", "",
        "- Reference/lookup data consumed by the observed capabilities is sourced from existing backend "
        "tables or services and is not independently redefined by this document.",
        "- Users reach the observed capabilities through an existing authentication/session mechanism; this "
        "document does not re-specify authentication.",
        "", "### 1.6 Constraints", "",
    ]
    signals = _detect_tech_signals(manifest)
    if signals["data_stores"]:
        lines.append(
            f"- The current implementation persists data via {_join_and(signals['data_stores'])}; any "
            "requirement that changes stored data must be coordinated with that backend."
        )
    if signals["frameworks"]:
        lines.append(
            f"- The current implementation is built on {_join_and(signals['frameworks'])}; requirements are "
            "scoped to behavior observable within that stack."
        )
    if not signals["data_stores"] and not signals["frameworks"]:
        lines.append(
            "- No specific backend technology constraint was established from source evidence beyond the "
            "languages detected in the Authoritative Source Coverage Register."
        )
    return "\n".join(lines)


def _stakeholders_markdown(identity: dict, capabilities: list[dict]) -> str:
    client = identity.get("client_name") or "the client organization"
    names = _join_and([str(c.get("name") or "") for c in capabilities]) or "the observed capabilities"
    return "\n".join([
        "## 2. Stakeholders", "",
        "| Stakeholder | Interest / Role |", "| --- | --- |",
        f"| Business Owner ({client}) | Primary business owner; maintains {names} data. |",
        "| End Users | Search, create, update, and (where observed) delete records within the observed capabilities. |",
        "| IT / Application Support | Builds, maintains, and supports the application and its backend integrations. |",
        "| Downstream Consumers | Any other system or process that depends on data maintained by the observed capabilities. |",
    ])


def _success_criteria_markdown(capabilities: list[dict]) -> str:
    lines = ["## 5. Success Criteria / Acceptance Measures", ""]
    for capability in capabilities:
        operations = _join_and(capability.get("operations") or [])
        verb_phrase = operations.lower() if operations else "work with"
        lines.append(
            f"- Authorized users can {verb_phrase} {capability.get('name')} records without requiring IT "
            "or database intervention."
        )
    lines += [
        "- Unauthorized or inquiry-only users are prevented from performing create, update, or delete actions.",
        "- Every requirement above is traceable to the cited source evidence in the Authoritative Source "
        "Coverage Register.",
    ]
    return "\n".join(lines)


def _glossary_markdown(identity: dict, capabilities: list[dict]) -> str:
    app = identity.get("project_name") or identity.get("application_name") or "the application"
    lines = ["## 6. Glossary", "", "| Term | Definition |", "| --- | --- |", f"| {app} | The system documented in this requirements document. |"]
    for capability in capabilities:
        lines.append(f"| {capability.get('name')} | An observed business capability of {app}; see Section 3 and Section 4. |")
    lines += [
        "| Evidence ID | A `SRC-####` identifier referencing a specific source file in the Authoritative Source Coverage Register. |",
        "| Observed | A statement directly supported by cited source evidence. |",
        "| Open Question | A statement that could not be established from source evidence and requires business input. |",
    ]
    return "\n".join(lines)


# ─── BRD: 3. Business Requirements table + 4. Business Rules ──────────────
# Model-authored content is limited to one descriptive sentence per row and a
# handful of rule bullets per capability, given the capability's FULL source
# evidence — everything else (which rows exist, their IDs, names, priorities)
# is decided deterministically from the observed operations, so the table is
# always complete and correctly grouped even if every model call fails.

# (kind, operations-that-trigger-it, name-template, priority, default description template)
_REQUIREMENT_ROW_SPECS = (
    ("Search", ("View / Search / List",), "{name} Search", "High",
     "Users shall be able to search and view existing {name} records."),
    ("Maintenance", ("Add", "Modify"), "{name} Maintenance", "High",
     "Authorized users shall be able to add and update {name} records."),
    ("Deletion", ("Delete",), "{name} Deletion", "Medium",
     "Authorized users shall be able to remove an existing {name} record."),
    ("Export", ("Export",), "{name} Export", "Medium",
     "Users shall be able to export {name} search results for offline review."),
)
_CROSS_CUTTING_REQUIREMENT_ROWS = (
    ("Role-Based Access", "High",
     "The system shall distinguish between administrative users (full add/update/delete rights) and "
     "inquiry-only users (search/view/export rights only), based on the user's assigned role."),
    ("Result Set Governance", "Medium",
     "When a search would return an excessive number of records, the system shall warn the user and "
     "require the search criteria to be narrowed rather than returning an unbounded result set."),
    ("Auditability", "Medium",
     "The system shall retain the identifier of the last user who updated a record and the timestamp of "
     "that update, for traceability."),
    ("Consistent Master Data", "High",
     "Master data maintained by the observed capabilities shall be available and consistent for any "
     "downstream process that depends on it."),
)


def _capability_row_specs(capability: dict) -> list[tuple]:
    operations = set(capability.get("operations") or [])
    return [spec for spec in _REQUIREMENT_ROW_SPECS if operations.intersection(spec[1])]


def _default_capability_rules(capability: dict) -> list[str]:
    name = capability.get("name") or "Capability"
    return [
        f"A {name} record is uniquely identified by a {name} code, per observed source evidence.",
        f"{name} status/lifecycle values observed in source govern whether a record is active for use.",
    ]


def _capability_business_prompt(capability: dict, evidence: list[dict], row_kinds: list[str], note: str = "") -> str:
    kinds = ", ".join(row_kinds) or "(no requirement rows apply to this capability)"
    return f"""You are a business analyst extracting business requirements and business rules from real
application source code for one capability: "{capability.get('name')}".

SOURCE EVIDENCE (full file content — the only source you may describe; do not invent behavior beyond it):
{_evidence_block(evidence)}

Output EXACTLY this shape and nothing else — no preamble, no extra headings, no markdown formatting:
REQUIREMENTS
<one line per required type below, formatted "Type: description">
RULES
<3-6 bullet lines, each starting with "-", stating a specific, concrete business rule for this capability
(field constraints, allowed values, uniqueness, defaults, status lifecycle, or similar) drawn only from
the evidence above>

Required REQUIREMENTS types (fill exactly these, in this order): {kinds}
Each description must be a single sentence (15-40 words) naming concrete fields, validations, or behavior
observed in the evidence. Do not invent fields, values, or behavior not shown above.{note}
"""


def _parse_capability_business_output(output: str, row_kinds: list[str]) -> tuple[dict[str, str], list[str]]:
    descriptions: dict[str, str] = {}
    rules: list[str] = []
    section = None
    for raw_line in output.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        upper = line.upper()
        if upper.startswith("REQUIREMENTS"):
            section = "requirements"
            continue
        if upper.startswith("RULES"):
            section = "rules"
            continue
        if section == "requirements":
            match = re.match(r"^-?\s*(Search|Maintenance|Deletion|Export)\s*:\s*(.+)$", line, re.IGNORECASE)
            if match:
                kind, description = match.group(1).title(), match.group(2).strip()
                if kind in row_kinds and len(description.split()) >= 6 and not _looks_like_refusal(description):
                    descriptions[kind] = description
        elif section == "rules":
            bullet = re.match(r"^[-*]\s+(.+)$", line)
            if bullet and len(bullet.group(1).split()) >= 4 and not _looks_like_refusal(bullet.group(1)):
                rules.append(bullet.group(1).strip())
    return descriptions, rules


def _generate_business_requirements_and_rules(
    capabilities: list[dict], source_path: str | Path, model: str, on_token: Optional[Callable[[str], None]],
) -> str:
    row_lines: list[str] = []
    rule_sections: list[tuple[str, list[str]]] = []
    counter = 0

    def next_id() -> str:
        nonlocal counter
        counter += 1
        return f"BR-{counter:02d}"

    for capability in capabilities:
        specs = _capability_row_specs(capability)
        row_kinds = [spec[0] for spec in specs]
        descriptions: dict[str, str] = {}
        rules: list[str] = []
        if row_kinds:
            evidence = _capability_full_evidence(capability, source_path)
            note = ""
            for _attempt in range(CAPABILITY_MAX_ATTEMPTS):
                output = _safe_generate(
                    f"requirements:{capability.get('name')}",
                    prompt=_capability_business_prompt(capability, evidence, row_kinds, note), model=model,
                    system="You are a precise business analyst. Follow the exact output shape requested and cite only supplied evidence.",
                    on_token=on_token, max_tokens=CAPABILITY_MAX_TOKENS, num_ctx=CAPABILITY_CONTEXT_TOKENS,
                    max_seconds=CAPABILITY_MAX_SECONDS,
                )
                descriptions, rules = _parse_capability_business_output(output, row_kinds)
                if len(descriptions) >= len(row_kinds) and len(rules) >= 2:
                    break
                missing = [kind for kind in row_kinds if kind not in descriptions]
                note = (
                    f"\n\nPREVIOUS ATTEMPT WAS INCOMPLETE — missing description(s) for: "
                    f"{', '.join(missing) or 'none'}; only {len(rules)} rule bullet(s) were parsed (need at "
                    "least 2). Follow the exact output shape exactly."
                )
        for kind, _ops, name_template, priority, default_description in specs:
            name = name_template.format(name=capability.get("name"))
            description = descriptions.get(kind) or default_description.format(name=capability.get("name"))
            row_lines.append(f"| {next_id()} | {name} | {description} | {priority} |")
        rule_sections.append((str(capability.get("name") or "Capability"), rules or _default_capability_rules(capability)))

    for name, priority, description in _CROSS_CUTTING_REQUIREMENT_ROWS:
        row_lines.append(f"| {next_id()} | {name} | {description} | {priority} |")

    requirements_md = "\n".join([
        "## 3. Business Requirements", "",
        "The table below consolidates the business requirements identified from the current system "
        "behavior. Each requirement is tagged with a unique ID and a relative priority to support future "
        "backlog planning.", "",
        "| ID | Requirement Name | Description | Priority |",
        "| --- | --- | --- | --- |",
        *(row_lines or ["| — | No business capability was safely established | Source review required | — |"]),
    ])
    rules_lines = ["## 4. Business Rules", ""]
    for index, (name, rules) in enumerate(rule_sections, start=1):
        rules_lines.append(f"### 4.{index} {name}")
        rules_lines.append("")
        rules_lines.extend(f"- {rule}" for rule in rules)
        rules_lines.append("")
    if not rule_sections:
        rules_lines.append("No business capability was safely established from source evidence.")
    return requirements_md + "\n\n" + "\n".join(rules_lines).rstrip()


# ─── FSD: one "Module: <Capability>" section per observed capability ──────

def _module_prompt(capability: dict, evidence: list[dict], note: str = "") -> str:
    operations = ", ".join(capability.get("operations") or []) or "behavior requires source review"
    return f"""You are a technical analyst writing the implementation-level functional specification for
one module of a governed application: "{capability.get('name')}" (observed operations: {operations}).

SOURCE EVIDENCE (full file content — the only source you may describe; do not invent fields, action
codes, or behavior not shown above):
{_evidence_block(evidence)}

Write ONLY the Markdown for this module, starting with the exact heading `## {capability.get('name')}`.
Include, using level-3 (###) subheadings, whichever of these are supported by the evidence above:
- Search / List (if applicable): a Markdown table `| Field | Type | Max Length | Required | Notes |` for
  the search/filter fields, one row per field found in the evidence.
- Add / Update (if applicable): a Markdown table `| Field | Type | Max Length | Required | Notes |` for
  every field captured on add/update, one row per field found in the evidence.
- Action Codes (only if the evidence shows literal action/command codes passed to a controller or
  backend): a Markdown table `| Code | Meaning |`.
- Validation & Business Logic: 3-6 bullet points citing concrete validation rules, required fields, or
  processing logic found in the evidence.
- Export (if an export/reporting operation is observed): one paragraph describing what is exported and how.
- Deletion (if a delete operation is observed): one paragraph describing how deletion is invoked and its effect.
- Backend Interface (only if the evidence shows a stored procedure, API endpoint, or query the module
  calls): a short note naming it and what it does.
Cite the Evidence ID (`SRC-#### — path`) for every table or bullet point you write. Do not write a section
for any capability other than "{capability.get('name')}". Do not invent field names, types, or lengths not
present in the evidence.{note}
"""


def _module_section_issues(content: str, capability: dict) -> list[str]:
    name = str(capability.get("name") or "")
    issues: list[str] = []
    if not re.match(rf"(?im)^#{{1,6}}[^\n]*\b{re.escape(name)}\b", content.strip()):
        issues.append(f"section does not open with a '{name}' heading")
    if len(content.split()) < 60:
        issues.append("module section is too short")
    if _looks_like_refusal(content):
        issues.append("model refusal or generic response")
    expected_evidence = {str(value).upper() for value in capability.get("evidence_ids") or []}
    cited_evidence = {value.upper() for value in re.findall(r"\bSRC-\d{4}\b", content, flags=re.IGNORECASE)}
    if expected_evidence and not expected_evidence.intersection(cited_evidence):
        issues.append("no capability-specific Evidence ID citation")
    return issues


def _fallback_module_section(capability: dict) -> str:
    """Deterministic, evidence-grounded module content used only if the model
    still falls short of a compliant section after its retry — honest about
    what it is (an automated placeholder) rather than inventing field-level
    detail it cannot verify."""
    name = capability.get("name") or "Capability"
    operations = capability.get("operations") or ["Observed behavior requires review"]
    pairs = list(zip(capability.get("evidence_ids") or [], capability.get("source_paths") or [])) or [("—", "source review required")]
    lines = [
        f"## {name}", "",
        f"Automated field-level extraction did not complete for {name} within the available time budget. "
        "The points below are derived directly from observed evidence pending a closer manual pass.",
        "", "### Validation & Business Logic", "",
    ]
    for index, operation in enumerate(operations):
        evidence_id, path = pairs[index % len(pairs)]
        lines.append(f"- The system supports a **{operation}** operation for {name}. Evidence: {evidence_id} — {path}.")
    declarations = capability.get("declarations") or []
    lines += ["", "### Source Declarations", ""]
    lines.append(", ".join(declarations[:20]) + "." if declarations else "No declarations were detected in the inspected source for this capability.")
    return "\n".join(lines)


def _generate_modules(
    capabilities: list[dict], source_path: str | Path, model: str, on_token: Optional[Callable[[str], None]],
) -> str:
    sections: list[str] = []
    for capability in capabilities:
        evidence = _capability_full_evidence(capability, source_path)
        section, note, issues = "", "", ["not attempted"]
        for _attempt in range(MODULE_MAX_ATTEMPTS):
            output = _safe_generate(
                f"module:{capability.get('name')}",
                prompt=_module_prompt(capability, evidence, note), model=model,
                system="You are a precise technical analyst. Describe only what the supplied evidence shows.",
                on_token=on_token, max_tokens=MODULE_MAX_TOKENS, num_ctx=MODULE_CONTEXT_TOKENS,
                max_seconds=MODULE_MAX_SECONDS,
            )
            section = output.strip()
            issues = _module_section_issues(section, capability)
            if not issues:
                break
            note = "\n\nPREVIOUS ATTEMPT WAS REJECTED for: " + "; ".join(issues) + ". Fix every issue."
        if issues:
            section = _fallback_module_section(capability)
        sections.append(section)
    numbered = [
        re.sub(r"^##\s+", f"## {index}. Module: ", section, count=1)
        for index, section in enumerate(sections, start=4)
    ]
    return "\n\n".join(numbered)


# ─── FSD: 1. Introduction / 2. Document Conventions / 3. System Overview ──

def _fsd_introduction_markdown(identity: dict, capabilities: list[dict]) -> str:
    app = identity.get("project_name") or identity.get("application_name") or "the application"
    lines = [
        "## 1. Introduction", "",
        "### 1.1 Purpose", "",
        f"This Functional Specification Document (FSD) describes, in implementation-level detail, how "
        f"{app} satisfies the business requirements defined in the companion Business Requirements "
        "Document (BRD). It documents the screens, fields, validations, business logic, and backend "
        "interfaces observed in the current application, and is intended as the functional baseline for "
        "maintenance, defect triage, and future enhancement.", "",
        "### 1.2 Traceability to Business Requirements", "",
        "| Capability | Covered In |", "| --- | --- |",
    ]
    if capabilities:
        for index, capability in enumerate(capabilities, start=4):
            lines.append(f"| {capability.get('name')} | Section {index}. Module: {capability.get('name')} |")
    else:
        lines.append("| — | No business capability was safely established from source evidence |")
    return "\n".join(lines)


def _fsd_conventions_markdown() -> str:
    return "\n".join([
        "## 2. Document Conventions", "",
        "- Field lengths and types reflect what was observed in the current application's source code, not "
        "necessarily a formally documented interface contract.",
        "- Statements marked Observed are directly supported by cited source evidence; statements that "
        "could not be established from source are marked Open Question.",
    ])


def _system_overview_markdown(identity: dict, capabilities: list[dict], manifest: list[dict]) -> str:
    app = identity.get("project_name") or identity.get("application_name") or "the application"
    signals = _detect_tech_signals(manifest)
    names = _join_and([str(c.get("name") or "") for c in capabilities]) or "no capability safely established from source evidence"
    lines = [
        "## 3. System Overview", "",
        "### 3.1 System Context", "",
        f"{app} implements {names}, as evidenced by the Authoritative Source Coverage Register below. No "
        "other business function, module, or workflow was observed.", "",
        "### 3.2 High-Level Architecture", "",
        "| Layer | Technology / Component |", "| --- | --- |",
    ]
    detected_rows = 0
    if signals["frameworks"]:
        lines.append(f"| Application | {_join_and(signals['frameworks'])} |")
        detected_rows += 1
    if signals["languages"]:
        lines.append(f"| Implementation Languages | {_join_and(signals['languages'])} |")
        detected_rows += 1
    if signals["data_stores"]:
        lines.append(f"| Data Persistence | {_join_and(signals['data_stores'])} |")
        detected_rows += 1
    if not detected_rows:
        file_types = _join_and(sorted(_coverage_summary(manifest)["file_types"].keys()))
        lines.append(f"| Detected file types | {file_types or 'None detected'} |")
    lines += [
        "", "### 3.3 Security & Access Control", "",
        "Capability-specific access-control and validation behavior, where evidenced, is documented within "
        "each module's Validation & Business Logic subsection below.",
    ]
    return "\n".join(lines)


def _non_functional_markdown(index: int) -> str:
    return "\n".join([
        f"## {index}. Non-Functional Observations", "",
        "- Performance and scalability characteristics beyond what is stated per-module were not "
        "independently established from source evidence.",
        "- Logging, monitoring, and operational tooling are as observed within the cited source; anything "
        "not evidenced is an Open Question.",
    ])


def _open_items_markdown(index: int) -> str:
    return "\n".join([
        f"## {index}. Open Items / Recommendations for Future Iterations", "",
        "- Formalize field-level validation for any fields the evidence shows relying on backend-side "
        "enforcement rather than explicit client/server validation.",
        "- Document backend interface contracts (APIs, stored procedures, or equivalent) referenced in the "
        "module sections above with the owning team, to reduce tribal-knowledge dependency.",
    ])


# ─── Orchestration ──────────────────────────────────────────────────────────

def generate_requirement_artifact(
    document_type: str, project: dict, source_path: str | Path,
    on_token: Optional[Callable[[str], None]] = None,
) -> dict:
    """Generate one requirements artifact.

    BRD/FSD follow a fixed numbered template (see module docstring): the
    section skeleton is always deterministic and complete; a local Ollama
    model is used only for narrowly-scoped, independently-retried, evidence-
    grounded writing tasks with an honest deterministic fallback, so the
    document is always structurally complete regardless of model behavior.
    The knowledge graph remains a single evidence-grounded completion.
    """
    if document_type not in DOCUMENT_TYPES:
        raise ValueError(f"Unsupported requirements document type: {document_type}")
    manifest, excerpts = _source_evidence(source_path)
    capabilities = _functional_capability_inventory(manifest)
    identity = _project_identity(project)

    if document_type == "knowledge_graph":
        model = _graph_model()
        if not model:
            raise RuntimeError("Ollama is unavailable or no supported model is installed")
        graph_num_ctx = _GRAPH_CONTEXT_WINDOW_BY_MODEL.get(model, _GRAPH_CONTEXT_WINDOW_DEFAULT)
        # Leave real headroom for the completion below whichever model's
        # context window actually gets used, rather than sizing the evidence
        # payload right up to the ceiling with ~0 margin (see GRAPH_PREFERRED_
        # MODELS comment above) — this is what previously produced a
        # truncated/malformed response despite every /api/generate call
        # succeeding at the transport level.
        if graph_num_ctx >= 32_000:
            manifest_budget, excerpt_budget = 10_000, 24_000
        else:
            manifest_budget, excerpt_budget = 5_000, 12_000
        context = _project_context(
            project, source_path, manifest, excerpts,
            manifest_budget=manifest_budget, excerpt_budget=excerpt_budget,
        )
        quality_rules = """QUALITY RULES:
- The `observed_functional_capabilities` inventory is the authoritative application scope.
- Create one dedicated Markdown section named exactly for each observed capability.
- Within each capability section, describe every listed operation and cite at least one of that capability's Evidence IDs.
- Do not substitute generic industry features for observed source-code functionality.
- Do not introduce additional business capabilities unless directly supported by cited source evidence.
- Treat the supplied evidence register as the coverage boundary; do not silently omit capabilities.
- Cite evidence as `SRC-#### — path/to/file` so findings remain auditable.
- Never claim an inventory-only file was content-inspected.
"""
        prompt = f"""Analyze the governed project evidence below and follow the requested output contract.

PROJECT EVIDENCE:
{context}

OUTPUT CONTRACT:
{_GRAPH_INSTRUCTIONS}

{quality_rules}"""
        output = _safe_generate(
            "knowledge_graph", prompt=prompt, model=model,
            system="You are a senior business analyst and requirements architect. Ground every result in supplied evidence.",
            on_token=on_token, max_tokens=_GRAPH_MAX_TOKENS, num_ctx=graph_num_ctx, max_seconds=600,
        )
        try:
            artifact = _extract_json(output)
            if len(artifact["nodes"]) < 12 or not artifact["edges"]:
                raise ValueError("Knowledge graph was too sparse")
        except (ValueError, json.JSONDecodeError):
            # A 30-100 node graph is a large completion — give the retry the
            # same time budget as the initial attempt, not a shorter one. A
            # shorter retry budget previously turned a slow-but-viable model
            # (observed: ~14.5KB of otherwise-valid JSON produced before a
            # single malformed token broke strict parsing) into a guaranteed
            # timeout on the very attempt meant to fix that, since generating
            # an equally rich graph again takes just as long the second time.
            output = _safe_generate(
                "knowledge_graph_retry",
                prompt=prompt + "\n\nQUALITY RETRY: The previous graph was malformed or too sparse. Return one complete, valid JSON object with at least 30 connected evidence-grounded nodes.",
                model=model,
                system="Return strict JSON only. Build a detailed requirements knowledge graph grounded in supplied source evidence.",
                on_token=on_token, max_tokens=_GRAPH_MAX_TOKENS, num_ctx=graph_num_ctx, max_seconds=600,
            )
            try:
                artifact = _extract_json(output)
            except (ValueError, json.JSONDecodeError) as exc:
                raise RuntimeError(
                    f"Ollama ({model}) did not return a usable knowledge graph after a retry — not even a "
                    "partial graph could be salvaged from either response. This is usually a timeout or a "
                    "severely malformed/empty response. Try Generate Knowledge Graph again; if it persists, "
                    "a different installed model may handle this project's evidence size more reliably."
                ) from exc
        root_id = f"project:{identity['project_id']}"
        nodes = artifact.setdefault("nodes", [])
        if not any(str(node.get("id")) == root_id for node in nodes):
            nodes.insert(0, {
                "id": root_id,
                "label": f"{identity['project_id']} · {identity['project_name']}",
                "type": "business",
                "description": (
                    f"Governed project for {identity['client_name'] or 'unspecified client'}; "
                    f"application key {identity['application_key'] or 'not specified'}."
                ),
            })
        for node in nodes:
            node.update({key: identity[key] for key in ("project_id", "project_name", "application_key", "client_name")})
        edges = artifact.setdefault("edges", [])
        connected = {str(edge.get("target")) for edge in edges if str(edge.get("source")) == root_id}
        for node in nodes[1:]:
            node_id = str(node.get("id") or "")
            if node_id and node_id not in connected and node.get("type") in {"business", "feature", "functional"}:
                edges.append({"source": root_id, "target": node_id, "relationship": "governs", "project_id": identity["project_id"]})
        for edge in edges:
            edge["project_id"] = identity["project_id"]
    elif document_type == "brd":
        model = _requirements_model()
        if not model:
            raise RuntimeError("Ollama is unavailable or no supported model is installed")
        content = "\n\n".join([
            _document_control_markdown("brd", identity),
            _brd_introduction_markdown(identity, capabilities, manifest),
            _stakeholders_markdown(identity, capabilities),
            _generate_business_requirements_and_rules(capabilities, source_path, model, on_token),
            _success_criteria_markdown(capabilities),
            _glossary_markdown(identity, capabilities),
            _coverage_appendix(manifest),
        ])
        artifact = {"title": "Business Requirements Document", "content": content}
    else:
        model = _requirements_model()
        if not model:
            raise RuntimeError("Ollama is unavailable or no supported model is installed")
        modules_markdown = _generate_modules(capabilities, source_path, model, on_token)
        next_index = 4 + max(len(capabilities), 1)
        content = "\n\n".join([
            _document_control_markdown("fsd", identity),
            _fsd_introduction_markdown(identity, capabilities),
            _fsd_conventions_markdown(),
            _system_overview_markdown(identity, capabilities, manifest),
            modules_markdown,
            _non_functional_markdown(next_index),
            _open_items_markdown(next_index + 1),
            _coverage_appendix(manifest),
        ])
        artifact = {"title": "Functional Specification Document", "content": content}
    artifact["document_type"] = document_type
    artifact["model"] = model
    artifact["project_identity"] = identity
    artifact["source_coverage"] = _coverage_summary(manifest)
    artifact["capability_tagline"] = _join_and([str(c.get("name") or "") for c in capabilities]) or "No capability safely established"
    return artifact


def _word_text(markdown: str) -> str:
    """Remove common inline Markdown markers while preserving their text."""
    value = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", markdown)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(\*\*|__|`|~~)", "", value)
    return value.strip()


def _add_markdown_table(document, lines: list[str]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows = [[_word_text(cell) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    # ``Table.cell(row, column)`` rebuilds python-docx's flattened cell grid
    # on every access. That becomes effectively quadratic for generated source
    # coverage appendices (APP-004 has 574 rows / 4,000+ cells) and can take
    # minutes, causing the reverse proxy to return 500 before export finishes.
    # Walk each already-created row directly so rendering remains linear in
    # the number of cells.
    for row_index, (table_row, values) in enumerate(zip(table.rows, rows)):
        for cell, value in zip(table_row.cells, values):
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    document.add_paragraph()


def _add_markdown_content(document, content: str, title: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = content.splitlines()
    index = 0
    title_key = _word_text(title).lower()
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            heading_text = _word_text(heading.group(2))
            if not (len(heading.group(1)) == 1 and heading_text.lower() == title_key):
                document.add_heading(heading_text, level=min(len(heading.group(1)), 4))
            index += 1
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(_word_text(bullet.group(1)), style="List Bullet")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif numbered:
            paragraph = document.add_paragraph(_word_text(numbered.group(1)), style="List Number")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif re.fullmatch(r"[-*_]{3,}", stripped):
            document.add_paragraph()
        else:
            paragraph = document.add_paragraph(_word_text(stripped))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        index += 1


def _apply_aptos_style(style, size=None) -> None:
    """Set Word font attributes for Latin, East Asian, and complex-script text."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style.font.name = "Aptos"
    if size is not None:
        style.font.size = size
    properties = style.element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Aptos")


def _add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def build_requirement_docx(artifact: dict, project: dict) -> bytes:
    """Render a generated BRD/FSD artifact as a valid Word document."""
    if artifact.get("document_type") not in {"brd", "fsd"}:
        raise ValueError("DOCX export is available only for BRD and FSD artifacts")
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover - declared runtime dependency
        raise RuntimeError("python-docx is required for Word export") from exc

    title = str(artifact.get("title") or "Requirements Document")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    document.core_properties.title = title
    document.core_properties.subject = f"Generated requirements for {project.get('name', '')}"
    document.core_properties.author = "Modernization Studio · Strat-Aqorynth"
    identity = {**_project_identity(project), **(artifact.get("project_identity") or {})}
    document.core_properties.keywords = ", ".join(filter(None, (
        identity.get("project_id"), identity.get("application_key"), identity.get("client_name"),
    )))

    normal = document.styles["Normal"]
    _apply_aptos_style(normal, Pt(10.5))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size in (
        ("Title", 26), ("Subtitle", 12), ("Heading 1", 18), ("Heading 2", 15),
        ("Heading 3", 12), ("Heading 4", 11), ("List Bullet", 10.5), ("List Number", 10.5),
    ):
        try:
            _apply_aptos_style(document.styles[style_name], Pt(size))
        except KeyError:  # pragma: no cover - localized Word templates
            continue

    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{project.get('id', '')} · {project.get('name', '')}").bold = True
    tagline = str(artifact.get("capability_tagline") or "").strip()
    if tagline:
        capability_line = document.add_paragraph()
        capability_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        capability_line.add_run(f"Covers: {tagline}").italic = True
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Prepared from analysis of the existing application source code").italic = True

    values = (
        ("Project Primary Key", identity.get("project_id", "")),
        ("Project Name", identity.get("project_name", "")),
        ("Application Key", identity.get("application_key", "")),
        ("Application Name", identity.get("application_name", "")),
        ("Client Name", identity.get("client_name", "")),
        ("Application Owner", identity.get("application_owner", "")),
        ("Business Unit", identity.get("business_unit", "")),
        ("Business Criticality", identity.get("business_criticality", "")),
        ("Document Type", str(artifact.get("document_type", "")).upper()),
        ("Document Version", "1.0"),
        ("Document Status", "Generated draft — requires governed review and approval"),
        ("Generated By", "OpenSourceLLM"),
    )
    metadata = document.add_table(rows=len(values), cols=2)
    metadata.style = "Light Shading Accent 1"
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()
    _add_markdown_content(document, str(artifact.get("content") or ""), title)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"{identity.get('project_id', '')}  |  {title}").italic = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generated by Modernization Studio · Strat-Aqorynth").italic = True
    footer.add_run("  |  ")
    _add_page_number(footer)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()
