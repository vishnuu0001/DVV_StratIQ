import re
import time

from unittest.mock import patch

import pytest

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from services.requirements_documentation import (
    _capability_full_evidence, _capability_row_specs, _detect_tech_signals, _extract_json,
    _fallback_module_section, _functional_capability_inventory, _generate_business_requirements_and_rules,
    _generate_modules, _governed_analysis, _graph_model, _module_section_issues,
    _parse_capability_business_output, _project_context, _safe_generate, _source_evidence,
    build_requirement_docx, generate_requirement_artifact,
)


def test_extract_json_keeps_only_edges_with_declared_nodes():
    graph = _extract_json('''```json
    {"title":"Project", "nodes":[{"id":"need","label":"Need"},{"id":"feature","label":"Feature"}],
     "edges":[{"source":"need","target":"feature","relationship":"enables"},
              {"source":"missing","target":"feature","relationship":"enables"}]}
    ```''')
    assert len(graph["edges"]) == 1
    assert graph["edges"][0]["source"] == "need"


def test_extract_json_salvages_a_response_broken_by_one_malformed_node():
    """Reproduces the real observed failure: a large, mostly-valid response
    with a single syntax glitch (a missing comma partway through) must not
    be discarded wholesale — every other well-formed node/edge should
    survive, since re-running an expensive ~30-100 node completion from
    scratch over one bad token is wasteful and not always faster or better."""
    broken = (
        '{"nodes":['
        '{"id":"a","label":"A","type":"business"},'
        '{"id":"b" "label":"Broken, missing colon-comma"},'  # malformed: no comma after "b"
        '{"id":"c","label":"C","type":"feature"}'
        '],"edges":['
        '{"source":"a","target":"c","relationship":"enables"}'
        ']}'
    )
    graph = _extract_json(broken)
    assert {node["id"] for node in graph["nodes"]} == {"a", "c"}
    assert len(graph["edges"]) == 1


def test_extract_json_still_raises_when_nothing_can_be_salvaged():
    with pytest.raises(ValueError):
        _extract_json("not json at all, no nodes array present")


def test_rejects_unknown_document_type():
    with pytest.raises(ValueError, match="Unsupported"):
        generate_requirement_artifact("unknown", {}, ".")


def test_graph_model_prefers_larger_context_models_over_deepseek_coder():
    """deepseek-coder:6.7b is REQUIREMENTS_PREFERRED_MODELS' first choice (fine
    for the small, per-capability BRD/FSD calls) but has only a 16,384-token
    native context — too little headroom for the knowledge graph's much
    larger evidence payload. The graph path must prefer a bigger-context
    model when one is installed, rather than reusing that default."""
    with patch(
        "services.requirements_documentation.llm.check_status",
        return_value={"models": ["deepseek-coder:6.7b", "qwen3.5:9b"]},
    ):
        assert _graph_model() == "qwen3.5:9b"
    with patch(
        "services.requirements_documentation.llm.check_status",
        return_value={"models": ["deepseek-coder:6.7b"]},
    ):
        assert _graph_model() == "deepseek-coder:6.7b"  # falls back when it's all that's installed


def test_knowledge_graph_sizes_evidence_payload_to_the_selected_models_context(tmp_path):
    """Reproduces the reported failure's root cause: with only deepseek-coder
    (16,384-token native context) installed, the evidence payload must be
    small enough to leave real headroom for the completion, not sized right
    up to the ceiling with ~0 margin."""
    (tmp_path / "a.py").write_text("def x():\n    return 1\n", encoding="utf-8")
    captured = {}

    def fake_generate(prompt=None, num_ctx=None, max_tokens=None, **_kwargs):
        captured["num_ctx"] = num_ctx
        captured["prompt_len"] = len(prompt or "")
        return '{"nodes":[' + ",".join(f'{{"id":"n{i}","label":"N{i}","type":"feature"}}' for i in range(15)) + '],"edges":[{"source":"n0","target":"n1","relationship":"uses"}]}'

    with patch("services.requirements_documentation.llm.check_status", return_value={"models": ["deepseek-coder:6.7b"]}), \
         patch("services.requirements_documentation.llm.generate", side_effect=fake_generate):
        generate_requirement_artifact("knowledge_graph", {
            "id": "APP-005", "name": "Small", "configuration": {},
        }, tmp_path)

    assert captured["num_ctx"] == 16_384
    # Prompt (in chars) plus the completion's max_tokens (~4 chars/token) must
    # leave meaningful headroom under the model's context window, not consume
    # nearly all of it the way the pre-fix unbounded payload did.
    assert captured["prompt_len"] < (captured["num_ctx"] - 4_096) * 3


def test_knowledge_graph_raises_a_clear_error_when_both_attempts_are_unparseable(tmp_path):
    (tmp_path / "a.py").write_text("def x():\n    return 1\n", encoding="utf-8")
    with patch("services.requirements_documentation.llm.check_status", return_value={"models": ["deepseek-coder:6.7b"]}), \
         patch("services.requirements_documentation.llm.generate", return_value="not json at all"):
        with pytest.raises(RuntimeError, match="did not return a usable knowledge graph"):
            generate_requirement_artifact("knowledge_graph", {
                "id": "APP-005", "name": "Small", "configuration": {},
            }, tmp_path)


def test_build_requirement_docx_creates_native_word_structures():
    content = """## Document Control

## 4. Module: Widget Setup
### Validation & Business Logic
1. **FS-001:** The system shall export Word documents.

| ID | Requirement |
| --- | --- |
| FS-001 | Export DOCX |
"""
    raw = build_requirement_docx(
        {
            "document_type": "fsd", "title": "Functional Specification Document", "content": content,
            "model": "test-model", "capability_tagline": "Widget Setup",
        },
        {"id": "APP-004", "name": "Example Project", "configuration": {
            "application_key": "EXAMPLE", "client_name": "Contoso", "application_owner": "Jane Doe",
        }},
    )
    assert raw.startswith(b"PK")
    document = Document(BytesIO(raw))
    assert any("FS-001" in paragraph.text for paragraph in document.paragraphs)
    assert any("Export DOCX" in cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert any("Covers: Widget Setup" in paragraph.text for paragraph in document.paragraphs)
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    assert "Project Primary Key" in table_text and "APP-004" in table_text
    assert "Application Key" in table_text and "EXAMPLE" in table_text
    assert "Client Name" in table_text and "Contoso" in table_text
    assert document.styles["Normal"].font.name == "Aptos"
    fonts = document.styles["Normal"].element.get_or_add_rPr().rFonts
    assert fonts.get(qn("w:ascii")) == "Aptos"
    fs_paragraph = next(paragraph for paragraph in document.paragraphs if "FS-001" in paragraph.text)
    assert fs_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert "Document Status" in table_text


def test_build_requirement_docx_renders_large_coverage_table_without_timeout():
    header = "| Evidence | Path | Type | Bytes | Lines | Coverage | Notes |"
    separator = "| --- | --- | --- | --- | --- | --- | --- |"
    rows = [
        f"| SRC-{index:04d} | src/File{index}.java | Java | 100 | 10 | inspected | retained |"
        for index in range(1, 601)
    ]
    started = time.monotonic()
    raw = build_requirement_docx(
        {
            "document_type": "brd",
            "title": "Business Requirements Document",
            "content": "\n".join(["## Source Coverage", header, separator, *rows]),
        },
        {"id": "APP-004", "name": "Large Java Project", "configuration": {}},
    )
    elapsed = time.monotonic() - started

    assert raw.startswith(b"PK")
    assert elapsed < 15, f"large DOCX table rendering took {elapsed:.2f}s"


def test_project_context_covers_manifest_semantics_and_source_evidence(tmp_path):
    (tmp_path / "README.md").write_text("Order processing business workflow", encoding="utf-8")
    source = tmp_path / "OrderService.py"
    source.write_text("def submit_order(order_id):\n    return order_id\n", encoding="utf-8")
    context = _project_context({"id": "APP-001", "name": "Orders", "snapshots": []}, tmp_path)
    assert '"source_files_discovered": 2' in context
    assert '"path":"OrderService.py"' in context
    assert "submit_order" in context
    assert "Order processing business workflow" in context
    assert "SRC-0001" in context and "SRC-0002" in context


def test_governed_analysis_loads_snapshot_artifact(tmp_path):
    snapshot = tmp_path / "analysis" / "v001"
    snapshot.mkdir(parents=True)
    (snapshot / "artifact.json").write_text('{"capabilities":["Order capture"]}', encoding="utf-8")
    analysis = _governed_analysis({"snapshots": [{"kind": "analysis", "path": str(snapshot)}]})
    assert analysis == {"capabilities": ["Order capture"]}


def test_capability_inventory_requires_cross_layer_evidence_and_finds_crud(tmp_path):
    files = {
        "app/action/CarrierSetupAction.java": "class CarrierSetupAction { void addCarrier(){} void deleteCarrier(){} }",
        "app/services/CarrierSetupService.java": "class CarrierSetupService { void updateCarrier(){} void searchCarrierList(){} }",
        "app/action/LocationIndexAction.java": "class LocationIndexAction { void saveLocation(){} void deleteLocation(){} }",
        "app/services/LocationInformationService.java": "class LocationInformationService { void updateLocation(){} void getLocationDetails(){} }",
        "app/pages/footer.jsp": "<button>Export</button>",
    }
    for relative, content in files.items():
        path = tmp_path / relative
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
    manifest, _ = _source_evidence(tmp_path)
    capabilities = _functional_capability_inventory(manifest)
    assert [item["name"] for item in capabilities] == ["Carrier Setup", "Location Setup"]
    assert {"Add", "Modify", "Delete", "View / Search / List"}.issubset(capabilities[0]["operations"])
    assert not any(item["name"] == "Footer Setup" for item in capabilities)


def test_safe_generate_converts_a_timeout_into_an_empty_result():
    """A single slow/failed completion must not blow up the whole multi-call
    pipeline — _safe_generate is what lets each capability call fail
    independently without aborting already-validated work elsewhere."""
    with patch(
        "services.requirements_documentation.llm.generate",
        side_effect=TimeoutError("Ollama generation exceeded the 600s per-file budget"),
    ):
        assert _safe_generate("module", prompt="x", model="test-model") == ""


def test_capability_row_specs_map_operations_to_consolidated_row_kinds():
    """Mirrors the reference document's real grouping: Add+Modify collapse into
    one 'Maintenance' row, and Search/Deletion/Export are each their own row
    only when the underlying operation was actually observed."""
    carrier = {"operations": ["Add", "Modify", "Delete", "View / Search / List", "Export"]}
    kinds = [spec[0] for spec in _capability_row_specs(carrier)]
    assert kinds == ["Search", "Maintenance", "Deletion", "Export"]

    read_only = {"operations": ["View / Search / List"]}
    assert [spec[0] for spec in _capability_row_specs(read_only)] == ["Search"]


def test_parse_capability_business_output_extracts_requirements_and_rules():
    output = """REQUIREMENTS
Search: Users can search carrier records by code, name, and status.
Maintenance: Authorized users can add and update carrier records with contract details.
RULES
- A carrier is uniquely identified by a five-character carrier code.
- Carrier status includes Active, Inactive, and Blocked values.
"""
    descriptions, rules = _parse_capability_business_output(output, ["Search", "Maintenance"])
    assert descriptions["Search"].startswith("Users can search")
    assert descriptions["Maintenance"].startswith("Authorized users can add")
    assert len(rules) == 2


def test_parse_capability_business_output_ignores_refusals_and_short_lines():
    output = "REQUIREMENTS\nSearch: no\nRULES\n- ok\n"
    descriptions, rules = _parse_capability_business_output(output, ["Search"])
    assert descriptions == {}  # "no" is below the 6-word minimum
    assert rules == []  # "ok" is below the 4-word minimum


def test_capability_full_evidence_reads_full_file_content(tmp_path):
    source = tmp_path / "app" / "CarrierSetupAction.java"
    source.parent.mkdir(parents=True)
    body = "class CarrierSetupAction {\n" + "\n".join(f"    // line {i}" for i in range(50)) + "\n}\n"
    source.write_text(body, encoding="utf-8")
    capability = {"evidence_ids": ["SRC-0001"], "source_paths": ["app/CarrierSetupAction.java"]}
    evidence = _capability_full_evidence(capability, tmp_path)
    assert len(evidence) == 1
    assert evidence[0]["content"] == body.strip()
    assert "// line 49" in evidence[0]["content"]  # not truncated like the old excerpt path


def test_detect_tech_signals_finds_framework_and_data_store_keywords():
    manifest = [
        {"path": "app/action/CarrierSetupAction.java", "symbols": ["struts.ActionServlet"], "type": "Java", "coverage": "content-inspected"},
        {"path": "xml/DAO/SV77954.xml", "symbols": [], "type": "XML", "coverage": "content-inspected"},
    ]
    signals = _detect_tech_signals(manifest)
    assert "Apache Struts" in signals["frameworks"]
    assert "Java" in signals["languages"]


def _transportation_setup_manifest(write_files_to=None):
    """A manifest matching the shape of the originally reported project: two
    capabilities, Carrier Setup and Location Setup, each with real CRUD
    operations, so tests exercise the actual reported scenario. When
    `write_files_to` (a tmp_path) is given, real (small but non-empty) files
    are written at each capability's source paths — required for any test
    that exercises _capability_full_evidence / module or requirement
    generation for real, since a fake path with no backing file legitimately
    can't be read and would silently fall back for every capability, hiding
    whether the model-authored path (and its evidence-citation check) works."""
    capability_entries = [
        {"evidence_id": "SRC-0001", "path": "app/action/CarrierSetupAction.java", "type": "Java",
         "bytes": 100, "lines": 20, "symbols": ["CarrierSetupAction"], "coverage": "content-inspected",
         "capability_terms": ["Carrier"], "operations": ["Add", "Modify", "Delete", "View / Search / List", "Export"]},
        {"evidence_id": "SRC-0002", "path": "app/services/CarrierSetupService.java", "type": "Java",
         "bytes": 100, "lines": 20, "symbols": ["CarrierSetupService"], "coverage": "content-inspected",
         "capability_terms": ["Carrier"], "operations": ["Modify", "View / Search / List"]},
        {"evidence_id": "SRC-0003", "path": "app/action/LocationIndexAction.java", "type": "Java",
         "bytes": 100, "lines": 20, "symbols": ["LocationIndexAction"], "coverage": "content-inspected",
         "capability_terms": ["Location"], "operations": ["Add", "Modify", "Delete", "View / Search / List", "Export"]},
        {"evidence_id": "SRC-0004", "path": "app/services/LocationInformationService.java", "type": "Java",
         "bytes": 100, "lines": 20, "symbols": ["LocationInformationService"], "coverage": "content-inspected",
         "capability_terms": ["Location"], "operations": ["Modify", "View / Search / List"]},
    ]
    filler = [
        {"evidence_id": f"SRC-{9000 + i:04d}", "path": f"filler/File{i}.txt", "type": "Text",
         "bytes": 10, "lines": 1, "symbols": [], "coverage": "inventory-only"}
        for i in range(176)
    ]
    manifest = capability_entries + filler  # 180 files, matching the reported project's manifest size
    excerpts = [{"evidence_id": item["evidence_id"], "path": item["path"], "excerpt": item["symbols"][0]} for item in capability_entries]
    if write_files_to is not None:
        for item in capability_entries:
            file_path = write_files_to / item["path"]
            file_path.parent.mkdir(parents=True, exist_ok=True)
            file_path.write_text(f"class {item['symbols'][0]} {{ /* {item['path']} */ }}\n", encoding="utf-8")
    return manifest, excerpts


_REQUIREMENTS_RULES_RESPONSE = """REQUIREMENTS
Search: Users can search existing records by code, name, and status, individually or in combination.
Maintenance: Authorized users can add and update records, capturing identity, type, and contact details.
Deletion: Authorized users can remove an existing record via the observed delete action.
Export: Users can export the current search results to a workbook for offline review.
RULES
- A record is uniquely identified by a code field observed in the DTO.
- Status values observed include Active, Inactive, and Blocked with an effective date.
- Mailing and shipping addresses are tracked independently per observed fields.
"""


def _module_style_response(prompt: str) -> str:
    # Long enough to clear _module_section_issues' word-count floor for real —
    # a too-short mock here would silently exercise the deterministic fallback
    # instead of the model-authored path while still passing (fallback also
    # produces the right heading), masking whether the success path even works.
    match = re.search(r"starting with the exact heading `## (.+?)`", prompt)
    name = match.group(1) if match else "Capability"
    evidence_match = re.search(r"Evidence ID (SRC-\d{4})", prompt)
    evidence_id = evidence_match.group(1) if evidence_match else "SRC-0001"
    return (
        f"## {name}\n\n### Add / Update\n\n"
        "| Field | Type | Max Length | Required | Notes |\n| --- | --- | --- | --- | --- |\n"
        f"| Code | Text | 5 | Yes | Unique {name} identifier observed in the form. |\n"
        f"| Name | Text | 30 | Yes | Display name captured on the {name} add/update screen. |\n\n"
        "### Validation & Business Logic\n\n"
        f"- The system validates required fields before save, per the observed form fields. Evidence: {evidence_id} — file.\n"
        f"- Status values are constrained to a fixed set observed in source. Evidence: {evidence_id} — file.\n"
        f"- The record code is treated as the unique key and is read-only once created. Evidence: {evidence_id} — file.\n"
    )


def _combined_llm_side_effect(prompt=None, **_kwargs):
    if prompt and "REQUIREMENTS" in prompt and "RULES" in prompt:
        return _REQUIREMENTS_RULES_RESPONSE
    if prompt and "starting with the exact heading" in prompt:
        return _module_style_response(prompt)
    return ""


def test_generate_business_requirements_and_rules_produces_the_reference_shape(tmp_path):
    manifest, _ = _transportation_setup_manifest(write_files_to=tmp_path)
    capabilities = _functional_capability_inventory(manifest)
    with patch("services.requirements_documentation.llm.generate", side_effect=_combined_llm_side_effect):
        markdown = _generate_business_requirements_and_rules(capabilities, tmp_path, "test-model", None)
    assert "## 3. Business Requirements" in markdown
    assert "## 4. Business Rules" in markdown
    assert "### 4.1 Carrier Setup" in markdown
    assert "### 4.2 Location Setup" in markdown
    # Carrier Setup: Search, Maintenance, Deletion, Export = 4 rows; Location
    # Setup the same = 4 rows; plus 4 fixed cross-cutting rows = 12 total —
    # realistic, not an inflated count tied to the 180-file manifest size.
    ids = set(re.findall(r"\bBR-\d{2}\b", markdown))
    assert len(ids) == 12
    assert "Role-Based Access" in markdown and "Consistent Master Data" in markdown


def test_generate_business_requirements_and_rules_falls_back_when_model_never_complies(tmp_path):
    """Even a model that never produces a parseable response must not cause a
    capability's requirement rows or business rules to be dropped."""
    manifest, _ = _transportation_setup_manifest()
    capabilities = _functional_capability_inventory(manifest)
    with patch("services.requirements_documentation.llm.generate", return_value="not a compliant response"):
        markdown = _generate_business_requirements_and_rules(capabilities, tmp_path, "test-model", None)
    assert "Carrier Setup Search" in markdown
    assert "Carrier Setup Maintenance" in markdown
    assert "A Carrier Setup record is uniquely identified" in markdown
    ids = set(re.findall(r"\bBR-\d{2}\b", markdown))
    assert len(ids) == 12  # rows still exist with templated descriptions


def test_module_section_issues_flags_missing_heading_and_evidence():
    capability = {"name": "Carrier Setup", "evidence_ids": ["SRC-0001"]}
    assert "no capability-specific Evidence ID citation" in _module_section_issues(
        "## Carrier Setup\n\nSome content with no citation at all here to speak of.", capability,
    )
    assert any("heading" in issue for issue in _module_section_issues("## Location Setup\n\nWrong heading.", capability))


def test_fallback_module_section_is_evidence_grounded():
    capability = {
        "name": "Carrier Setup", "operations": ["Add", "Delete"],
        "evidence_ids": ["SRC-0001"], "source_paths": ["CarrierSetupAction.java"], "declarations": ["CarrierSetupAction"],
    }
    section = _fallback_module_section(capability)
    assert section.startswith("## Carrier Setup")
    assert "SRC-0001" in section
    assert not _module_section_issues(section, capability)


def test_generate_modules_falls_back_when_model_never_complies(tmp_path):
    manifest, _ = _transportation_setup_manifest()
    capabilities = _functional_capability_inventory(manifest)
    with patch("services.requirements_documentation.llm.generate", return_value="not a compliant module"):
        markdown = _generate_modules(capabilities, tmp_path, "test-model", None)
    assert "## 4. Module: Carrier Setup" in markdown
    assert "## 5. Module: Location Setup" in markdown


def test_generate_requirement_artifact_produces_the_reference_brd_template(tmp_path):
    manifest, excerpts = _transportation_setup_manifest(write_files_to=tmp_path)
    with patch("services.requirements_documentation._source_evidence", return_value=(manifest, excerpts)), \
         patch("services.requirements_documentation._requirements_model", return_value="test-model"), \
         patch("services.requirements_documentation.llm.generate", side_effect=_combined_llm_side_effect):
        artifact = generate_requirement_artifact("brd", {
            "id": "APP-003", "name": "TransportationSetup",
            "configuration": {"application_key": "TRANSPORT", "client_name": "Mazda"},
        }, tmp_path)

    content = artifact["content"]
    assert artifact["document_type"] == "brd"
    assert artifact["title"] == "Business Requirements Document"
    assert artifact["capability_tagline"] == "Carrier Setup and Location Setup"
    for heading in (
        "## Document Control", "### Version History", "### Distribution List",
        "## 1. Introduction", "### 1.1 Purpose", "### 1.2 Background", "### 1.3 Business Objectives",
        "### 1.4 Scope", "#### 1.4.1 In Scope", "#### 1.4.2 Out of Scope",
        "### 1.5 Assumptions", "### 1.6 Constraints", "## 2. Stakeholders",
        "## 3. Business Requirements", "## 4. Business Rules", "## 5. Success Criteria / Acceptance Measures",
        "## 6. Glossary", "## Appendix A: Observed Functional Capability Register",
        "## Appendix B: Authoritative Source Coverage Register",
    ):
        assert heading in content, f"missing {heading!r}"
    assert "Carrier Setup" in content and "Location Setup" in content
    # Content accuracy: never generic industry boilerplate instead of the named capabilities.
    assert "comprehensive solution" not in content.lower()


def test_generate_requirement_artifact_produces_the_reference_fsd_template(tmp_path):
    manifest, excerpts = _transportation_setup_manifest(write_files_to=tmp_path)
    with patch("services.requirements_documentation._source_evidence", return_value=(manifest, excerpts)), \
         patch("services.requirements_documentation._requirements_model", return_value="test-model"), \
         patch("services.requirements_documentation.llm.generate", side_effect=_combined_llm_side_effect):
        artifact = generate_requirement_artifact("fsd", {
            "id": "APP-003", "name": "TransportationSetup",
            "configuration": {"application_key": "TRANSPORT", "client_name": "Mazda"},
        }, tmp_path)

    content = artifact["content"]
    assert artifact["document_type"] == "fsd"
    assert artifact["title"] == "Functional Specification Document"
    for heading in (
        "## Document Control", "## 1. Introduction", "### 1.1 Purpose",
        "### 1.2 Traceability to Business Requirements", "## 2. Document Conventions",
        "## 3. System Overview", "### 3.1 System Context", "### 3.2 High-Level Architecture",
        "### 3.3 Security & Access Control", "## 4. Module: Carrier Setup", "## 5. Module: Location Setup",
        "## 6. Non-Functional Observations", "## 7. Open Items / Recommendations for Future Iterations",
        "## Appendix A: Observed Functional Capability Register",
    ):
        assert heading in content, f"missing {heading!r}"
    # Prove the model-authored path actually landed, not the deterministic
    # fallback (which also produces the headings above but not a field table
    # or this exact wording) — a too-short mock could otherwise pass this test
    # while silently only exercising the fallback.
    assert "Add / Update" in content
    assert "Unique Carrier Setup identifier observed in the form." in content
    assert "Unique Location Setup identifier observed in the form." in content
    assert "Automated field-level extraction did not complete" not in content


def test_generate_requirement_artifact_survives_every_ollama_call_timing_out():
    """The hardened failure mode: Ollama is reachable (a model is configured)
    but every single completion times out. Generation must still produce a
    complete, template-compliant document via the deterministic fallbacks
    rather than raising — every section is either fully deterministic or has
    a deterministic fallback, so there is no remaining failure path."""
    manifest, excerpts = _transportation_setup_manifest()
    with patch("services.requirements_documentation._source_evidence", return_value=(manifest, excerpts)), \
         patch("services.requirements_documentation._requirements_model", return_value="test-model"), \
         patch(
             "services.requirements_documentation.llm.generate",
             side_effect=TimeoutError("Ollama generation exceeded the 600s per-file budget"),
         ):
        brd = generate_requirement_artifact("brd", {
            "id": "APP-003", "name": "TransportationSetup",
            "configuration": {"application_key": "TRANSPORT", "client_name": "Mazda"},
        }, ".")
        fsd = generate_requirement_artifact("fsd", {
            "id": "APP-003", "name": "TransportationSetup",
            "configuration": {"application_key": "TRANSPORT", "client_name": "Mazda"},
        }, ".")

    assert "Carrier Setup" in brd["content"] and "Location Setup" in brd["content"]
    assert "## 3. Business Requirements" in brd["content"] and "## 4. Business Rules" in brd["content"]
    assert "## 4. Module: Carrier Setup" in fsd["content"] and "## 5. Module: Location Setup" in fsd["content"]
